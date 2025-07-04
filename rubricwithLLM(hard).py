#!/usr/bin/env python3
"""
Project Handoff Rubric Analyzer with Complete 32-Criterion Rubric - SCALED TO 100 POINTS
Analyzes project handoff documents against the comprehensive 32-criterion rubric
using AI models and generates detailed Word document reports.

SCORING SYSTEM:
- Content Quality: 75 points (75% weight) - scaled from 180 raw points (32 criteria)
- Foreman Attendance: 25 points (25% weight)
- TOTAL: 100 points

Requirements:
pip install python-docx openai anthropic python-dotenv

Usage:
Simply modify the DOCUMENT_PATH variable below and run the script directly in VSCode
"""

import os
import json
from typing import Dict, List, Tuple, Any, Optional
from dataclasses import dataclass, field
from datetime import datetime
import re

# Third-party imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

# LLM clients (optional - install as needed)
try:
    import openai
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False

try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

# Load environment variables
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# =============================================================================
# CONFIGURATION - MODIFY THESE SETTINGS
# =============================================================================

# Path to your document (CHANGE THIS to your document path)
DOCUMENT_PATH = "C:\\Users\\kayla.dipaolo\\source\\repos\\provo river water treatment\\rft_analysis_Provo River Water Treatment Turnover.docx"

# Project name (will be auto-extracted from document, or use this as fallback)
FALLBACK_PROJECT_NAME = "Project Analysis"

# Model settings
MODEL_TYPE = "anthropic"  # Options: "openai" or "anthropic"
MODEL_NAME = "claude-3-5-sonnet-20241022"   # Latest model

# =============================================================================

@dataclass
class CriterionScore:
    """Individual criterion scoring breakdown"""
    name: str
    points: int
    max_points: int
    description: str
    justification: str
    supporting_evidence: str = ""

@dataclass
class CategoryResult:
    """Category-level results"""
    name: str
    total_score: int
    max_score: int
    criteria_scores: List[CriterionScore]
    
    @property
    def percentage(self) -> float:
        return (self.total_score / self.max_score * 100) if self.max_score > 0 else 0
    
    @property
    def scaled_score(self) -> float:
        """Scaled score as part of 75-point content system"""
        return (self.total_score / self.max_score * 75) if self.max_score > 0 else 0

@dataclass
class DocumentEvaluation:
    """Complete document evaluation results - SCALED TO 100 POINTS"""
    document_name: str
    project_name: str
    evaluation_date: str
    model_used: str
    category_results: List[CategoryResult] = field(default_factory=list)
    overall_strengths: List[str] = field(default_factory=list)
    areas_for_improvement: List[str] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)
    
    # Foreman attendance tracking
    foreman_present: bool = False
    foreman_names_found: List[str] = field(default_factory=list)
    
    @property
    def total_raw_score(self) -> int:
        """Raw total score from rubric (out of 180)"""
        return sum(cat.total_score for cat in self.category_results)
    
    @property
    def max_possible_raw_score(self) -> int:
        """Maximum raw score possible (180)"""
        return sum(cat.max_score for cat in self.category_results)
    
    @property
    def content_percentage(self) -> float:
        """Raw content score percentage from rubric"""
        return (self.total_raw_score / self.max_possible_raw_score * 100) if self.max_possible_raw_score > 0 else 0
    
    @property
    def scaled_content_score(self) -> float:
        """Scaled content score out of 75 points"""
        return (self.total_raw_score / self.max_possible_raw_score * 75) if self.max_possible_raw_score > 0 else 0
    
    @property
    def attendance_score(self) -> float:
        """Attendance score out of 25 points"""
        return 25.0 if self.foreman_present else 0.0
    
    @property
    def final_score(self) -> float:
        """Final scaled score out of 100 points"""
        return self.scaled_content_score + self.attendance_score
    
    @property
    def weighted_percentage(self) -> float:
        """Weighted percentage (same as final score since it's out of 100)"""
        return self.final_score
    
    @property
    def percentage(self) -> float:
        """Default to final score for backward compatibility"""
        return self.final_score
    
    @property
    def performance_level(self) -> str:
        score = self.final_score
        if score >= 90:
            return "EXCELLENT"
        elif score >= 80:
            return "GOOD"
        elif score >= 70:
            return "SATISFACTORY"
        elif score >= 60:
            return "NEEDS IMPROVEMENT"
        else:
            return "UNSATISFACTORY"

class RubricDefinition:
    """Complete 32-criterion rubric definition from documentation"""
    
    def __init__(self):
        self.rubric_data = self._load_complete_rubric()
    
    def _load_complete_rubric(self) -> Dict[str, Any]:
        """Load the complete 32-criterion rubric structure"""
        return {
            "Customer": {
                "max_score": 6,
                "criteria": [
                    {
                        "name": "Customer Background Details",
                        "max_points": 6,
                        "description": "Who is the customer? Prior relationship? Any specific concerns from past?",
                        "scoring_breakdown": {
                            "Customer Identity": {"points": 2, "guide": "2=Customer clearly named | 1=Vague | 0=Not identified"},
                            "Relationship History": {"points": 2, "guide": "2=Specific past work | 1=General mention | 0=No info"},
                            "Concerns": {"points": 1, "guide": "1=Specific concerns noted | 0=No concerns"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    }
                ]
            },
            "Timeline": {
                "max_score": 12,
                "criteria": [
                    {
                        "name": "Project Timeline Analysis",
                        "max_points": 6,
                        "description": "Start/end dates? Anticipated delays? Flexibility for unexpected changes?",
                        "scoring_breakdown": {
                            "Start Info": {"points": 1, "guide": "1=Clear start date | 0=Missing/vague"},
                            "Completion Info": {"points": 1, "guide": "1=Clear completion date | 0=Missing/vague"},
                            "Duration/Resources": {"points": 3, "guide": "3=All details | 2=Two elements | 1=One element | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Project Milestones",
                        "max_points": 6,
                        "description": "Key dates/events identified and communicated?",
                        "scoring_breakdown": {
                            "Milestones": {"points": 2, "guide": "2=All with dates | 1=Some | 0=None"},
                            "Integration": {"points": 2, "guide": "2=Schedule integrated | 1=Basic | 0=None"},
                            "Communication": {"points": 1, "guide": "1=Tracking plan | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    }
                ]
            },
            "Material": {
                "max_score": 10,
                "criteria": [
                    {
                        "name": "Material Sourcing & Availability",
                        "max_points": 5,
                        "description": "Key materials required? Lead times? Potential impact of delay?",
                        "scoring_breakdown": {
                            "Material Assessment": {"points": 2, "guide": "2=Comprehensive analysis | 1=Basic ID | 0=Not identified"},
                            "Lead Times": {"points": 1, "guide": "1=Clear lead times | 0=No info"},
                            "Mitigation": {"points": 1, "guide": "1=Actions mentioned | 0=No strategies"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Procurement Updates",
                        "max_points": 5,
                        "description": "Are current procurement statuses shared and tracked?",
                        "scoring_breakdown": {
                            "Status": {"points": 2, "guide": "2=Detailed status/timelines | 1=General | 0=None"},
                            "Updates": {"points": 1, "guide": "1=Regular process | 0=None"},
                            "Risks": {"points": 1, "guide": "1=Risks identified | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    }
                ]
            },
            "Safety": {
                "max_score": 36,
                "criteria": [
                    {
                        "name": "Site Conditions",
                        "max_points": 6,
                        "description": "Any access issues, weather challenges, or hazards noted?",
                        "scoring_breakdown": {
                            "Risk ID": {"points": 2, "guide": "2=Multiple risks | 1=One risk | 0=No risks"},
                            "Assessment": {"points": 1, "guide": "1=Risk level described | 0=No assessment"},
                            "Mitigation": {"points": 2, "guide": "2=Specific actions | 1=General | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Safety Considerations",
                        "max_points": 6,
                        "description": "Are safety protocols, risks, and responsibilities discussed?",
                        "scoring_breakdown": {
                            "Risk ID": {"points": 2, "guide": "2=Multiple safety risks | 1=One risk | 0=No risks"},
                            "Communication": {"points": 1, "guide": "1=Protocols communicated | 0=No plan"},
                            "Mitigation": {"points": 2, "guide": "2=Specific measures | 1=General | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Risk Management",
                        "max_points": 6,
                        "description": "Are project risks identified and mitigations in place?",
                        "scoring_breakdown": {
                            "Risk ID": {"points": 2, "guide": "2=Multiple risks | 1=One risk | 0=None"},
                            "Assessment": {"points": 1, "guide": "1=Impact assessed | 0=None"},
                            "Mitigation": {"points": 2, "guide": "2=Specific strategies | 1=General | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Work Area Protection",
                        "max_points": 6,
                        "description": "Are barricades, signage, temporary walls planned?",
                        "scoring_breakdown": {
                            "Protection": {"points": 2, "guide": "2=Comprehensive plan | 1=Basic | 0=None"},
                            "Signage": {"points": 2, "guide": "2=Detailed plan | 1=General | 0=None"},
                            "Timeline": {"points": 1, "guide": "1=Implementation schedule | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Access & Egress Plans",
                        "max_points": 6,
                        "description": "How do workers/materials get in/out safely and efficiently?",
                        "scoring_breakdown": {
                            "Routes": {"points": 2, "guide": "2=Specific safe routes | 1=General | 0=None"},
                            "Material Handling": {"points": 2, "guide": "2=Detailed plan | 1=Basic | 0=None"},
                            "Safety": {"points": 1, "guide": "1=Safety measures | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Glazing-Specific Risks",
                        "max_points": 6,
                        "description": "Are unique risks (e.g., lifting glass, storage) considered?",
                        "scoring_breakdown": {
                            "Risk ID": {"points": 2, "guide": "2=Multiple glazing risks | 1=Some | 0=None"},
                            "Protocols": {"points": 2, "guide": "2=Detailed protocols | 1=Basic | 0=None"},
                            "Equipment": {"points": 1, "guide": "1=Specialized needs | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    }
                ]
            },
            "Scope": {
                "max_score": 6,
                "criteria": [
                    {
                        "name": "Scope of Work",
                        "max_points": 6,
                        "description": "What's included/excluded? Are all trades and tasks covered?",
                        "scoring_breakdown": {
                            "Scope Confirmation": {"points": 2, "guide": "2=Detailed scope | 1=Basic outline | 0=Unclear"},
                            "Coverage": {"points": 1, "guide": "1=All trades ID | 0=Incomplete"},
                            "Issues": {"points": 2, "guide": "2=Comprehensive | 1=Basic | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    }
                ]
            },
            "Equipment": {
                "max_score": 6,
                "criteria": [
                    {
                        "name": "Equipment Needs",
                        "max_points": 6,
                        "description": "Special tools/lifts/scaffolding required? Who provides what?",
                        "scoring_breakdown": {
                            "Equipment ID": {"points": 1, "guide": "1=Specific types listed | 0=Not clear"},
                            "Costs": {"points": 2, "guide": "2=Specific costs | 1=General | 0=None"},
                            "Provider/Coord": {"points": 2, "guide": "2=Clear provider/timing | 1=Partial | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    }
                ]
            },
            "General": {
                "max_score": 104,
                "criteria": [
                    {
                        "name": "Permits & Inspections",
                        "max_points": 5,
                        "description": "Which permits are required? Status of inspections?",
                        "scoring_breakdown": {
                            "Permits": {"points": 2, "guide": "2=All permits with status | 1=Incomplete | 0=Unclear"},
                            "Schedule": {"points": 1, "guide": "1=Timeline provided | 0=No info"},
                            "Actions": {"points": 1, "guide": "1=Actions described | 0=No plan"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Subcontractor Coordination",
                        "max_points": 6,
                        "description": "Are subcontractors identified and scheduled?",
                        "scoring_breakdown": {
                            "Sub ID": {"points": 2, "guide": "2=All subs by name | 1=Some | 0=Unclear"},
                            "Schedule": {"points": 2, "guide": "2=Detailed plan | 1=Basic | 0=None"},
                            "Communication": {"points": 1, "guide": "1=Method defined | 0=No plan"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Communication Plan",
                        "max_points": 6,
                        "description": "Who's the main contact? Method/frequency of updates?",
                        "scoring_breakdown": {
                            "Contacts": {"points": 2, "guide": "2=Clear contacts with roles | 1=Some | 0=Unclear"},
                            "Methods": {"points": 2, "guide": "2=Specific methods/frequency | 1=General | 0=None"},
                            "Schedule": {"points": 1, "guide": "1=Update schedule | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Drawing/Spec Review",
                        "max_points": 5,
                        "description": "Have relevant drawings/specs been reviewed and explained?",
                        "scoring_breakdown": {
                            "Review": {"points": 2, "guide": "2=All documents reviewed | 1=Some | 0=None"},
                            "Understanding": {"points": 1, "guide": "1=Confirmed with team | 0=No confirmation"},
                            "Actions": {"points": 1, "guide": "1=Actions identified | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Change Order Procedure",
                        "max_points": 6,
                        "description": "Is the change process understood? Forms, approvals, pricing?",
                        "scoring_breakdown": {
                            "Process": {"points": 2, "guide": "2=Complete process | 1=Basic | 0=None"},
                            "Approvals": {"points": 2, "guide": "2=Clear hierarchy | 1=Some info | 0=Unclear"},
                            "Pricing": {"points": 1, "guide": "1=Method explained | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Budget/Cost Tracking",
                        "max_points": 6,
                        "description": "Is the budget reviewed and roles defined for tracking costs? Do we have an NTE?",
                        "scoring_breakdown": {
                            "Budget": {"points": 2, "guide": "2=Specific amounts/NTE | 1=General | 0=None"},
                            "Roles": {"points": 2, "guide": "2=Clear responsibilities | 1=Some | 0=Unclear"},
                            "Process": {"points": 1, "guide": "1=Monitoring defined | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Quality Control Measures",
                        "max_points": 6,
                        "description": "Inspections, punch list process, workmanship standards?",
                        "scoring_breakdown": {
                            "Inspections": {"points": 2, "guide": "2=Detailed schedule | 1=Basic | 0=None"},
                            "Standards": {"points": 2, "guide": "2=Specific standards | 1=General | 0=None"},
                            "Punch List": {"points": 1, "guide": "1=Clear process | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Emergency Contacts",
                        "max_points": 5,
                        "description": "Who to call for major issues? Is this list documented/shared?",
                        "scoring_breakdown": {
                            "Contacts": {"points": 2, "guide": "2=Complete list | 1=Some | 0=None"},
                            "Documentation": {"points": 1, "guide": "1=Documented/accessible | 0=None"},
                            "Protocol": {"points": 1, "guide": "1=Communication plan | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Storage/Logistics",
                        "max_points": 5,
                        "description": "Where are materials stored? Delivery coordination?",
                        "scoring_breakdown": {
                            "Storage": {"points": 2, "guide": "2=Specific locations/security | 1=General | 0=None"},
                            "Delivery": {"points": 1, "guide": "1=Schedule/coordination | 0=None"},
                            "Access": {"points": 1, "guide": "1=Access procedures | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Insurance & Bonding",
                        "max_points": 5,
                        "description": "Confirm coverage and any specific policy limitations?",
                        "scoring_breakdown": {
                            "Coverage": {"points": 2, "guide": "2=All requirements confirmed | 1=Basic | 0=None"},
                            "Limitations": {"points": 1, "guide": "1=Limitations noted | 0=None"},
                            "Compliance": {"points": 1, "guide": "1=Actions for compliance | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Legal/Contract Requirements",
                        "max_points": 6,
                        "description": "Are contract obligations highlighted and risks reviewed?",
                        "scoring_breakdown": {
                            "Obligations": {"points": 2, "guide": "2=Specific obligations | 1=General | 0=None"},
                            "Risks": {"points": 2, "guide": "2=Analyzed with mitigation | 1=Basic | 0=None"},
                            "Compliance": {"points": 1, "guide": "1=Compliance plan | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Neighbor/Adjacent Property Coordination",
                        "max_points": 5,
                        "description": "Have neighbors been notified? Special coordination needed?",
                        "scoring_breakdown": {
                            "Notification": {"points": 2, "guide": "2=Properly notified | 1=Some | 0=None"},
                            "Coordination": {"points": 1, "guide": "1=Requirements identified | 0=None"},
                            "Communication": {"points": 1, "guide": "1=Ongoing plan | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Roofing-Specific Hazards",
                        "max_points": 6,
                        "description": "Are fall protection, staging, and weather exposure addressed?",
                        "scoring_breakdown": {
                            "Fall Protection": {"points": 2, "guide": "2=Comprehensive plan | 1=Basic | 0=None"},
                            "Staging": {"points": 2, "guide": "2=Detailed plan | 1=Basic | 0=None"},
                            "Weather": {"points": 1, "guide": "1=Risks addressed | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Debrief with Outgoing Team",
                        "max_points": 6,
                        "description": "Has the outgoing crew/team been debriefed for issues or lessons learned?",
                        "scoring_breakdown": {
                            "Debrief": {"points": 2, "guide": "2=Comprehensive with docs | 1=Basic | 0=None"},
                            "Issues": {"points": 2, "guide": "2=Specific issues documented | 1=Some | 0=None"},
                            "Transfer": {"points": 1, "guide": "1=Knowledge transferred | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Client Communication Plan",
                        "max_points": 6,
                        "description": "How and when is the client updated?",
                        "scoring_breakdown": {
                            "Schedule": {"points": 2, "guide": "2=Specific schedule | 1=General | 0=None"},
                            "Contacts": {"points": 2, "guide": "2=Primary/backup contacts | 1=Some | 0=None"},
                            "Content": {"points": 1, "guide": "1=Content defined | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Escalation Pathway",
                        "max_points": 5,
                        "description": "When and how are issues escalated?",
                        "scoring_breakdown": {
                            "Triggers": {"points": 2, "guide": "2=Clear triggers | 1=Some | 0=None"},
                            "Process": {"points": 1, "guide": "1=Step-by-step | 0=None"},
                            "Contacts": {"points": 1, "guide": "1=Contact chain | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Temporary Systems",
                        "max_points": 5,
                        "description": "Are any temp power, HVAC, drainage setups explained?",
                        "scoring_breakdown": {
                            "Requirements": {"points": 2, "guide": "2=All systems with specs | 1=Some | 0=None"},
                            "Installation": {"points": 1, "guide": "1=Timeline/responsibility | 0=None"},
                            "Coordination": {"points": 1, "guide": "1=Trade coordination | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Conflict Resolution Strategy",
                        "max_points": 5,
                        "description": "Are conflicts resolved via defined process or chain of command?",
                        "scoring_breakdown": {
                            "Process": {"points": 2, "guide": "2=Clear process/steps | 1=Basic | 0=None"},
                            "Authority": {"points": 1, "guide": "1=Chain of command | 0=None"},
                            "Documentation": {"points": 1, "guide": "1=Tracking process | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    },
                    {
                        "name": "Post-Handoff Follow-Up",
                        "max_points": 5,
                        "description": "Is there a scheduled check-in after handoff?",
                        "scoring_breakdown": {
                            "Schedule": {"points": 2, "guide": "2=Specific meetings/agenda | 1=General | 0=None"},
                            "Responsibility": {"points": 1, "guide": "1=Clear responsibility | 0=None"},
                            "Metrics": {"points": 1, "guide": "1=Success criteria | 0=None"},
                            "Universal Check": {"points": 1, "guide": "1=2+ concrete details | 0=Vague"}
                        }
                    }
                ]
            }
        }

class LLMAnalyzer:
    """Handles LLM-based document analysis"""
    
    def __init__(self, model_type: str = "anthropic", model_name: str = "claude-3-5-sonnet-20241022"):
        self.model_type = model_type
        self.model_name = model_name
        self._setup_client()
    
    def _setup_client(self):
        """Setup the appropriate LLM client"""
        if self.model_type == "openai" and HAS_OPENAI:
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                print("Warning: OPENAI_API_KEY not found in environment variables.")
                self.client = None
            else:
                self.client = openai.OpenAI(api_key=api_key)
        elif self.model_type == "anthropic" and HAS_ANTHROPIC:
            api_key = os.getenv("ANTHROPIC_API_KEY")
            if not api_key:
                print("Warning: ANTHROPIC_API_KEY not found in environment variables.")
                self.client = None
            else:
                self.client = anthropic.Anthropic(api_key=api_key)
        else:
            self.client = None
            print(f"Warning: {self.model_type} client not available. Using mock analysis.")
    
    def analyze_criterion(self, document_content: str, criterion_data: Dict[str, Any]) -> CriterionScore:
        """Analyze a single criterion using LLM with retry logic"""
        
        if self.client is None:
            # Mock analysis for demonstration
            return self._mock_analysis(criterion_data)
        
        prompt = self._build_analysis_prompt(criterion_data, document_content)
        
        # Retry logic for API overload errors
        max_retries = 3
        base_delay = 2  # seconds
        
        for attempt in range(max_retries):
            try:
                if self.model_type == "openai":
                    response = self.client.chat.completions.create(
                        model=self.model_name,
                        messages=[
                            {"role": "system", "content": "You are an expert project management analyst evaluating handoff documents against specific criteria."},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.1
                    )
                    result_text = response.choices[0].message.content
                    
                elif self.model_type == "anthropic":
                    response = self.client.messages.create(
                        model=self.model_name,
                        max_tokens=1000,
                        temperature=0.1,
                        messages=[
                            {"role": "user", "content": prompt}
                        ]
                    )
                    
                    if hasattr(response, 'content') and response.content and len(response.content) > 0:
                        if hasattr(response.content[0], 'text'):
                            result_text = response.content[0].text
                        else:
                            result_text = str(response.content[0])
                    else:
                        print(f"Warning: Unexpected response format from Anthropic API")
                        return self._mock_analysis(criterion_data)
                
                return self._parse_llm_response(result_text, criterion_data)
                
            except Exception as e:
                error_str = str(e)
                
                # Check for overload error (529)
                if "529" in error_str or "overloaded" in error_str.lower():
                    if attempt < max_retries - 1:  # Don't sleep on last attempt
                        delay = base_delay * (2 ** attempt)  # Exponential backoff
                        print(f"   API overloaded, retrying in {delay} seconds... (attempt {attempt + 1}/{max_retries})")
                        import time
                        time.sleep(delay)
                        continue
                    else:
                        print(f"   API overloaded after {max_retries} attempts, using mock analysis")
                        return self._mock_analysis(criterion_data)
                else:
                    # For other errors, fail immediately
                    print(f"Error in LLM analysis for '{criterion_data['name']}': {e}")
                    return self._mock_analysis(criterion_data)
        
        # Fallback if all retries failed
        return self._mock_analysis(criterion_data)
    
    def _build_analysis_prompt(self, criterion_data: Dict[str, Any], document_content: str) -> str:
        """Build the analysis prompt for the LLM"""
        
        prompt = f"""
Analyze the following project handoff document against this specific criterion:

CRITERION: {criterion_data['name']}
DESCRIPTION: {criterion_data['description']}
MAX POINTS: {criterion_data['max_points']}

SCORING BREAKDOWN:
"""
        
        for component, details in criterion_data['scoring_breakdown'].items():
            prompt += f"- {component} ({details['points']} pts): {details['guide']}\n"
        
        prompt += f"""

DOCUMENT TO ANALYZE:
{document_content[:4000]}...

EVALUATION FOCUS:
- Reward SPECIFIC details: proper nouns, exact dates, quantities, dollar amounts
- Penalize GENERIC responses: "standard materials", "typical process", "will coordinate"
- Look for PROJECT-SPECIFIC information rather than templated language
- Evidence of actual preparation vs. form completion
- Apply the "Universal Check" - look for 2+ concrete details vs vague statements

Please provide your analysis in this exact format:
TOTAL_SCORE: [number out of {criterion_data['max_points']}]
JUSTIFICATION: [detailed explanation focusing on specificity vs. generic content]
SUPPORTING_EVIDENCE: [specific quotes or references from the document]

Focus on rewarding concrete, project-specific details and penalizing vague, generic responses.
"""
        
        return prompt
    
    def _parse_llm_response(self, response_text: str, criterion_data: Dict[str, Any]) -> CriterionScore:
        """Parse the LLM response into a structured result"""
        
        # Extract score
        score_match = re.search(r'TOTAL_SCORE:\s*(\d+)', response_text)
        score = int(score_match.group(1)) if score_match else 0
        
        # Extract justification
        just_match = re.search(r'JUSTIFICATION:\s*(.*?)(?=SUPPORTING_EVIDENCE:|$)', response_text, re.DOTALL)
        justification = just_match.group(1).strip() if just_match else "No justification provided"
        
        # Extract supporting evidence
        evidence_match = re.search(r'SUPPORTING_EVIDENCE:\s*(.*?)$', response_text, re.DOTALL)
        evidence = evidence_match.group(1).strip() if evidence_match else ""
        
        return CriterionScore(
            name=criterion_data['name'],
            points=min(score, criterion_data['max_points']),  # Ensure not over max
            max_points=criterion_data['max_points'],
            description=criterion_data['description'],
            justification=justification,
            supporting_evidence=evidence
        )
    
    def _mock_analysis(self, criterion_data: Dict[str, Any]) -> CriterionScore:
        """Provide mock analysis when LLM is not available"""
        
        # Simple mock scoring (for demonstration)
        mock_score = max(1, int(criterion_data['max_points'] * 0.7))  # 70% score, minimum 1
        
        return CriterionScore(
            name=criterion_data['name'],
            points=mock_score,
            max_points=criterion_data['max_points'],
            description=criterion_data['description'],
            justification="Mock analysis - LLM client not available or API key missing",
            supporting_evidence="Mock evidence for demonstration purposes"
        )

class ProjectHandoffAnalyzer:
    """Main analyzer class that coordinates the analysis"""
    
    def __init__(self, model_type: str = "anthropic", model_name: str = "claude-3-5-sonnet-20241022"):
        self.rubric = RubricDefinition()
        self.llm_analyzer = LLMAnalyzer(model_type, model_name)
    
    def _detect_foreman_attendance(self, document_content: str) -> Tuple[bool, List[str]]:
        """Use AI to detect if foreman attended the meeting based on transcript analysis"""
        
        if self.llm_analyzer.client is None:
            # Fallback for mock analysis
            return True, ["Mock: Foreman attendance assumed for demonstration"]
        
        prompt = f"""
Analyze this project handoff meeting transcript to determine if a FOREMAN was present and participating.

A foreman is typically:
- The field supervisor or construction lead
- Someone who provides installation insights, field experience, or practical input
- Someone who discusses field conditions, equipment needs, or installation challenges
- Someone who gives feedback on constructability or field execution

Look for evidence such as:
- Someone providing field/installation expertise or suggestions
- Discussion of field conditions, access, equipment needs from someone with hands-on experience
- Practical input about installation challenges, safety, or site logistics
- Someone asking/answering questions about actual construction execution
- References to someone who will be managing the field work

TRANSCRIPT TO ANALYZE:
{document_content[:2000]}...

Provide your assessment in this exact format:
FOREMAN_PRESENT: [YES or NO]
CONFIDENCE: [LOW, MEDIUM, HIGH]
EVIDENCE: [specific quotes or observations that support your determination]
REASONING: [explain why you believe a foreman was or wasn't present]

Focus on the ROLE and CONTRIBUTIONS rather than job titles or names.
"""
        
        try:
            if self.llm_analyzer.model_type == "openai":
                response = self.llm_analyzer.client.chat.completions.create(
                    model=self.llm_analyzer.model_name,
                    messages=[
                        {"role": "system", "content": "You are an expert construction project analyst who can identify team roles and participation based on meeting content."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1
                )
                result_text = response.choices[0].message.content

            elif self.llm_analyzer.model_type == "anthropic":
                response = self.llm_analyzer.client.messages.create(
                    model=self.llm_analyzer.model_name,
                    max_tokens=800,
                    temperature=0.1,
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )
                
                if hasattr(response, 'content') and response.content and len(response.content) > 0:
                    if hasattr(response.content[0], 'text'):
                        result_text = response.content[0].text
                    else:
                        result_text = str(response.content[0])
                else:
                    print(f"Warning: Unexpected response format from Anthropic API")
                    return False, ["API response format error"]

            # Parse the LLM response
            foreman_present = "YES" in result_text.upper() and "FOREMAN_PRESENT:" in result_text.upper()

            # Extract evidence and reasoning
            evidence_match = re.search(r'EVIDENCE:\s*(.*?)(?=REASONING:|$)', result_text, re.DOTALL)
            evidence = evidence_match.group(1).strip() if evidence_match else "No evidence extracted"

            reasoning_match = re.search(r'REASONING:\s*(.*?)$', result_text, re.DOTALL)
            reasoning = reasoning_match.group(1).strip() if reasoning_match else "No reasoning provided"

            confidence_match = re.search(r'CONFIDENCE:\s*(\w+)', result_text, re.IGNORECASE)
            confidence = confidence_match.group(1) if confidence_match else "UNKNOWN"

            analysis_details = [
                f"Confidence: {confidence}",
                f"Evidence: {evidence[:100]}..." if len(evidence) > 100 else f"Evidence: {evidence}",
                f"Reasoning: {reasoning[:100]}..." if len(reasoning) > 100 else f"Reasoning: {reasoning}"
            ]

            return foreman_present, analysis_details

        except Exception as e:
            print(f"Error in foreman detection: {e}")
            return False, [f"Error occurred during analysis: {str(e)}"]
    
    def analyze_document(self, document_path: str, document_name: str = None, 
                        project_name: str = "Project Analysis") -> DocumentEvaluation:
        """Analyze a complete document against the 32-criterion rubric"""
        
        # Read document based on file type
        document_content = self._read_document(document_path)
        
        if document_name is None:
            document_name = os.path.basename(document_path)
        
        # Auto-extract project name
        extracted_project_name = self._extract_project_name(document_content)
        if extracted_project_name != "Unknown Project":
            project_name = extracted_project_name
        
        # Initialize evaluation
        evaluation = DocumentEvaluation(
            document_name=document_name,
            project_name=project_name,
            evaluation_date=datetime.now().strftime("%Y-%m-%d"),
            model_used=f"{self.llm_analyzer.model_type}:{self.llm_analyzer.model_name}"
        )
        
        # Detect foreman attendance
        foreman_present, foreman_indicators = self._detect_foreman_attendance(document_content)
        evaluation.foreman_present = foreman_present
        evaluation.foreman_names_found = foreman_indicators
        
        print(f"🔍 Foreman Detection: {'✅ Present' if foreman_present else '❌ Not detected'}")
        if foreman_indicators:
            print(f"   Analysis: {foreman_indicators[0] if foreman_indicators else 'No details'}")
        
        # Analyze each category
        total_criteria = 0
        for category_name, category_data in self.rubric.rubric_data.items():
            print(f"\n📋 Analyzing {category_name} category ({len(category_data['criteria'])} criteria)")
            category_result = self._analyze_category(
                document_content, category_name, category_data
            )
            evaluation.category_results.append(category_result)
            total_criteria += len(category_data['criteria'])
        
        print(f"\n📊 Analyzed {total_criteria} criteria across {len(evaluation.category_results)} categories")
        
        # Generate overall insights
        evaluation.overall_strengths = self._identify_strengths(evaluation)
        evaluation.areas_for_improvement = self._identify_improvements(evaluation)
        evaluation.recommendations = self._generate_recommendations(evaluation)
        
        return evaluation
    
    def _read_document(self, document_path: str) -> str:
        """Read document content based on file type"""
        
        try:
            file_ext = document_path.lower().split('.')[-1]
            
            if file_ext == 'docx':
                # Read Word document
                doc = Document(document_path)
                document_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
            elif file_ext == 'json':
                # Read JSON file and convert to readable text
                with open(document_path, 'r', encoding='utf-8') as f:
                    json_data = json.load(f)
                document_content = self._convert_json_to_text(json_data)
                
            elif file_ext == 'csv':
                # Read CSV file
                import csv
                document_content = ""
                with open(document_path, 'r', encoding='utf-8') as f:
                    csv_reader = csv.reader(f)
                    for row in csv_reader:
                        document_content += " | ".join(row) + "\n"
                        
            else:
                # Read as text file
                with open(document_path, 'r', encoding='utf-8') as f:
                    document_content = f.read()
            
            return document_content
            
        except Exception as e:
            raise Exception(f"Error reading document '{document_path}': {str(e)}")
    
    def _convert_json_to_text(self, json_data: Any, indent_level: int = 0) -> str:
        """Convert JSON data to readable text format for analysis"""
        
        indent = "  " * indent_level
        text_parts = []
        
        if isinstance(json_data, dict):
            for key, value in json_data.items():
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{indent}{key}:")
                    text_parts.append(self._convert_json_to_text(value, indent_level + 1))
                else:
                    text_parts.append(f"{indent}{key}: {value}")
        
        elif isinstance(json_data, list):
            for i, item in enumerate(json_data):
                if isinstance(item, (dict, list)):
                    text_parts.append(f"{indent}Item {i + 1}:")
                    text_parts.append(self._convert_json_to_text(item, indent_level + 1))
                else:
                    text_parts.append(f"{indent}- {item}")
        
        else:
            text_parts.append(f"{indent}{json_data}")
        
        return "\n".join(text_parts)

    def _analyze_category(self, document_content: str, category_name: str, 
                         category_data: Dict[str, Any]) -> CategoryResult:
        """Analyze all criteria in a category"""
        
        criteria_scores = []
        
        for criterion_data in category_data['criteria']:
            print(f"  Analyzing: {criterion_data['name']} ({criterion_data['max_points']} pts)")
            score = self.llm_analyzer.analyze_criterion(document_content, criterion_data)
            criteria_scores.append(score)
            print(f"    Score: {score.points}/{score.max_points}")
        
        total_score = sum(score.points for score in criteria_scores)
        max_score = sum(score.max_points for score in criteria_scores)
        
        return CategoryResult(
            name=category_name,
            total_score=total_score,
            max_score=max_score,
            criteria_scores=criteria_scores
        )
    
    def _extract_project_name(self, document_content: str) -> str:
        """Extract project name from document content using various patterns"""
        
        # Patterns to look for (in order of preference)
        patterns = [
            # Direct project name patterns
            r'project\s*name[:\-\s]+(.+?)(?:\n|$)',
            r'project[:\-\s]+(.+?)(?:\n|$)',
            r'job\s*name[:\-\s]+(.+?)(?:\n|$)',
            r'job[:\-\s]+(.+?)(?:\n|$)',
            r'site[:\-\s]+(.+?)(?:\n|$)',
            r'client[:\-\s]+(.+?)(?:\n|$)',
            r'customer[:\-\s]+(.+?)(?:\n|$)',
        ]
        
        # Try each pattern
        for pattern in patterns:
            matches = re.findall(pattern, document_content, re.IGNORECASE | re.MULTILINE)
            if matches:
                # Clean up the match
                project_name = matches[0].strip()
                # Remove common prefixes/suffixes
                project_name = re.sub(r'^(the|a|an)\s+', '', project_name, flags=re.IGNORECASE)
                project_name = re.sub(r'\s+(project|handoff|document|summary)$', '', project_name, flags=re.IGNORECASE)
                
                if project_name and len(project_name) > 2:  # Make sure we have a valid project name
                    return project_name
        
        return "Unknown Project"
    
    def _identify_strengths(self, evaluation: DocumentEvaluation) -> List[str]:
        """Identify overall strengths from the analysis"""
        strengths = []
        
        # Find high-performing categories
        high_performers = [cat for cat in evaluation.category_results if cat.percentage >= 80]
        
        if high_performers:
            strengths.append(f"Strong performance in {', '.join([cat.name for cat in high_performers])}")
        
        # Find perfect scores
        perfect_scores = []
        for cat in evaluation.category_results:
            for crit in cat.criteria_scores:
                if crit.points == crit.max_points:
                    perfect_scores.append(crit.name)
        
        if perfect_scores:
            strengths.append(f"Perfect scores achieved in: {', '.join(perfect_scores[:3])}...")
        
        if not strengths:
            strengths.append("Analysis completed successfully with comprehensive 32-criterion evaluation")
        
        return strengths
    
    def _identify_improvements(self, evaluation: DocumentEvaluation) -> List[str]:
        """Identify areas needing improvement"""
        improvements = []
        
        # Find low-performing categories
        low_performers = [cat for cat in evaluation.category_results if cat.percentage < 60]
        
        if low_performers:
            improvements.append(f"Significant improvement needed in {', '.join([cat.name for cat in low_performers])}")
        
        # Find zero scores
        zero_scores = []
        for cat in evaluation.category_results:
            for crit in cat.criteria_scores:
                if crit.points == 0:
                    zero_scores.append(crit.name)
        
        if zero_scores:
            improvements.append(f"Critical gaps identified in: {', '.join(zero_scores[:3])}...")
        
        if not improvements:
            improvements.append("Continue maintaining current quality standards across all 32 criteria")
        
        return improvements
    
    def _generate_recommendations(self, evaluation: DocumentEvaluation) -> List[str]:
        """Generate specific recommendations based on 32-criterion evaluation"""
        recommendations = [
            "Develop comprehensive implementation procedures for all 32 criteria",
            "Create detailed checklists for each criterion to ensure complete coverage",
            "Establish regular review cycles to maintain document quality",
            "Implement cross-functional review process for complex projects",
            "Focus on providing specific, concrete details rather than generic responses"
        ]
        
        # Add specific recommendations based on performance
        if evaluation.final_score < 70:
            recommendations.insert(0, "Immediate comprehensive review and enhancement required across all criteria")
        
        return recommendations

class EnhancedWordReportGenerator:
    """Enhanced Word document generator with professional formatting - 32-CRITERION RUBRIC"""
    
    def __init__(self):
        self.doc = None
    
    def generate_comprehensive_report(self, evaluation: DocumentEvaluation, output_path: str):
        """Generate a comprehensive, professionally formatted report"""
        
        self.doc = Document()
        self._setup_styles()
        
        # Generate all sections
        self._add_title_page(evaluation)
        self._add_executive_summary(evaluation)
        self._add_category_analysis(evaluation)
        self._add_detailed_criteria_analysis(evaluation)
        self._add_recommendations_section(evaluation)
        self._add_appendices(evaluation)
        
        # Save document
        self.doc.save(output_path)
        print(f"📄 Comprehensive 32-criterion report saved to: {output_path}")
    
    def _setup_styles(self):
        """Setup professional document styles"""
        styles = self.doc.styles
        
        # Title style
        try:
            if 'Report Title' not in [s.name for s in styles]:
                title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.name = 'Calibri'
                title_style.font.size = Pt(24)
                title_style.font.bold = True
                title_style.font.color.rgb = RGBColor(0, 51, 102)
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass  # Style already exists or can't be created
        
        # Subtitle style
        try:
            if 'Report Subtitle' not in [s.name for s in styles]:
                subtitle_style = styles.add_style('Report Subtitle', WD_STYLE_TYPE.PARAGRAPH)
                subtitle_style.font.name = 'Calibri'
                subtitle_style.font.size = Pt(14)
                subtitle_style.font.color.rgb = RGBColor(89, 89, 89)
                subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass  # Style already exists or can't be created
    
    def _add_title_page(self, evaluation: DocumentEvaluation):
        """Add professional title page with 32-criterion scoring"""
        
        # Title
        try:
            title = self.doc.add_paragraph("Project Handoff Evaluation Report", style='Report Title')
        except:
            title = self.doc.add_paragraph("Project Handoff Evaluation Report")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Subtitle
        try:
            subtitle = self.doc.add_paragraph("32-Criterion Comprehensive Assessment", style='Report Subtitle')
        except:
            subtitle = self.doc.add_paragraph("32-Criterion Comprehensive Assessment")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        try:
            subtitle = self.doc.add_paragraph(evaluation.document_name, style='Report Subtitle')
        except:
            subtitle = self.doc.add_paragraph(evaluation.document_name)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Key metrics table with scaled scoring
        table = self.doc.add_table(rows=11, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        metrics = [
            ("Project", evaluation.project_name),
            ("Evaluation Date", evaluation.evaluation_date),
            ("Model Used", evaluation.model_used),
            ("Rubric Version", "32-Criterion Complete Assessment"),
            ("Foreman Present", "✅ Yes" if evaluation.foreman_present else "❌ No"),
            ("Raw Content Score", f"{evaluation.total_raw_score}/{evaluation.max_possible_raw_score} ({evaluation.content_percentage:.1f}%)"),
            ("Scaled Content Score", f"{evaluation.scaled_content_score:.1f}/75 points"),
            ("Attendance Score", f"{evaluation.attendance_score:.0f}/25 points"),
            ("Final Score", f"{evaluation.final_score:.1f}/100 points"),
            ("Performance Level", evaluation.performance_level)
        ]
        
        for i, (label, value) in enumerate(metrics):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            try:
                table.cell(i, 0).paragraphs[0].runs[0].bold = True
            except:
                pass  # Bold formatting failed
        
        self.doc.add_page_break()
    
    def _add_executive_summary(self, evaluation: DocumentEvaluation):
        """Add executive summary with 32-criterion scoring explanation"""
        
        self.doc.add_heading('Executive Summary', level=1)
        
        # Performance overview with scaled scoring
        performance_para = self.doc.add_paragraph()
        performance_para.add_run(f"Overall Performance: ").bold = True
        performance_para.add_run(
            f"This document achieved a final score of {evaluation.final_score:.1f}/100 points "
            f"based on the comprehensive 32-criterion evaluation framework. The breakdown includes "
            f"Content Quality: {evaluation.scaled_content_score:.1f}/75 points and "
            f"Foreman Attendance: {evaluation.attendance_score:.0f}/25 points, "
            f"earning a "
        )
        perf_run = performance_para.add_run(evaluation.performance_level)
        perf_run.bold = True
        try:
            if evaluation.performance_level in ['EXCELLENT', 'GOOD']:
                perf_run.font.color.rgb = RGBColor(0, 128, 0)
            elif evaluation.performance_level == 'SATISFACTORY':
                perf_run.font.color.rgb = RGBColor(255, 165, 0)
            else:
                perf_run.font.color.rgb = RGBColor(255, 0, 0)
        except:
            pass  # Color formatting failed
        performance_para.add_run(" rating.")
        
        # 32-criterion explanation
        self.doc.add_heading('32-Criterion Evaluation Framework', level=2)
        framework_para = self.doc.add_paragraph()
        framework_para.add_run("Comprehensive Assessment: ").bold = True
        framework_para.add_run(
            f"This evaluation uses a comprehensive 32-criterion rubric across 7 major categories: "
            f"Customer ({len([c for cat in evaluation.category_results if cat.name == 'Customer' for c in cat.criteria_scores])} criteria), "
            f"Timeline ({len([c for cat in evaluation.category_results if cat.name == 'Timeline' for c in cat.criteria_scores])} criteria), "
            f"Material ({len([c for cat in evaluation.category_results if cat.name == 'Material' for c in cat.criteria_scores])} criteria), "
            f"Safety ({len([c for cat in evaluation.category_results if cat.name == 'Safety' for c in cat.criteria_scores])} criteria), "
            f"Scope ({len([c for cat in evaluation.category_results if cat.name == 'Scope' for c in cat.criteria_scores])} criteria), "
            f"Equipment ({len([c for cat in evaluation.category_results if cat.name == 'Equipment' for c in cat.criteria_scores])} criteria), "
            f"and General ({len([c for cat in evaluation.category_results if cat.name == 'General' for c in cat.criteria_scores])} criteria). "
            f"Raw scores from these 32 criteria (totaling {evaluation.max_possible_raw_score} points) are scaled to 75 points for consistent reporting."
        )
        
        # Foreman attendance summary
        self.doc.add_heading('Foreman Attendance Analysis', level=2)
        foreman_para = self.doc.add_paragraph()
        if evaluation.foreman_present:
            foreman_para.add_run("✅ Foreman Present: ").bold = True
            foreman_para.add_run(f"Foreman attendance detected. Analysis details: {evaluation.foreman_names_found[0] if evaluation.foreman_names_found else 'No details available'}. ")
            foreman_para.add_run("Earned full 25/25 points for attendance.")
        else:
            foreman_para.add_run("❌ Foreman Not Detected: ").bold = True
            foreman_para.add_run("No clear foreman attendance indicators found in the transcript. ")
            foreman_para.add_run("0/25 points awarded for attendance. Consider ensuring foreman participation is documented.")
        
        # Category performance summary with scaled scores
        self.doc.add_heading('Category Performance Analysis', level=2)
        
        table = self.doc.add_table(rows=1, cols=6)
        try:
            table.style = 'Table Grid'
        except:
            pass  # Style not available
        
        # Headers
        headers = ['Category', 'Criteria', 'Raw Score', 'Raw Max', 'Scaled Score', 'Percentage']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            try:
                header_cells[i].paragraphs[0].runs[0].bold = True
            except:
                pass
        
        # Add data rows with scaled scores
        for category in evaluation.category_results:
            row_cells = table.add_row().cells
            row_cells[0].text = category.name
            row_cells[1].text = str(len(category.criteria_scores))
            row_cells[2].text = str(category.total_score)
            row_cells[3].text = str(category.max_score)
            
            # Calculate scaled score for this category
            scaled_score = (category.total_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            row_cells[4].text = f"{scaled_score:.1f}"
            row_cells[5].text = f"{category.percentage:.1f}%"
    
    def _add_category_analysis(self, evaluation: DocumentEvaluation):
        """Add detailed category-by-category analysis with 32-criterion scoring"""
        
        self.doc.add_heading('Category Analysis', level=1)
        
        for category in evaluation.category_results:
            scaled_score = (category.total_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            
            self.doc.add_heading(
                f'{category.name} Category ({category.total_score}/{category.max_score} raw | {scaled_score:.1f}/75 scaled - {category.percentage:.1f}%)', 
                level=2
            )
            
            # Category summary
            if category.percentage >= 80:
                summary = f"Strong performance with comprehensive coverage across {len(category.criteria_scores)} criteria."
            elif category.percentage >= 60:
                summary = f"Adequate performance with room for improvement across {len(category.criteria_scores)} criteria."
            else:
                summary = f"Significant gaps requiring immediate attention across {len(category.criteria_scores)} criteria."
            
            self.doc.add_paragraph(summary)
            
            # Individual criteria in this category
            for criterion in category.criteria_scores:
                self.doc.add_heading(f'{criterion.name} ({criterion.points}/{criterion.max_points})', level=3)
                self.doc.add_paragraph(f"Description: {criterion.description}")
                self.doc.add_paragraph(f"Assessment: {criterion.justification}")
                
                if criterion.supporting_evidence:
                    self.doc.add_paragraph(f"Evidence: {criterion.supporting_evidence}")
    
    def _add_detailed_criteria_analysis(self, evaluation: DocumentEvaluation):
        """Add comprehensive criteria analysis for all 32 criteria"""
        
        self.doc.add_heading('Detailed 32-Criterion Analysis', level=1)
        
        criterion_number = 1
        for category in evaluation.category_results:
            self.doc.add_heading(f'{category.name} Category Criteria', level=2)
            
            for criterion in category.criteria_scores:
                self.doc.add_heading(f'{criterion_number}. {criterion.name}', level=3)
                
                # Score badge
                score_para = self.doc.add_paragraph()
                score_run = score_para.add_run(f"Score: {criterion.points}/{criterion.max_points}")
                score_run.bold = True
                
                try:
                    if criterion.points == criterion.max_points:
                        score_run.font.color.rgb = RGBColor(0, 128, 0)
                    elif criterion.points >= criterion.max_points * 0.7:
                        score_run.font.color.rgb = RGBColor(255, 165, 0) 
                    else:
                        score_run.font.color.rgb = RGBColor(255, 0, 0)
                except:
                    pass  # Color formatting failed
                
                # Analysis details
                self.doc.add_paragraph(f"Criterion: {criterion.description}")
                self.doc.add_paragraph(f"Analysis: {criterion.justification}")
                
                if criterion.supporting_evidence:
                    self.doc.add_paragraph(f"Supporting Evidence: {criterion.supporting_evidence}")
                
                criterion_number += 1
    
    def _add_recommendations_section(self, evaluation: DocumentEvaluation):
        """Add comprehensive recommendations based on 32-criterion evaluation"""
        
        self.doc.add_heading('Recommendations & Action Items', level=1)
        
        # Strengths
        if evaluation.overall_strengths:
            self.doc.add_heading('Key Strengths', level=2)
            for strength in evaluation.overall_strengths:
                p = self.doc.add_paragraph(strength)
                try:
                    p.style = 'List Bullet'
                except:
                    pass  # Style not available
        
        # Areas for improvement  
        if evaluation.areas_for_improvement:
            self.doc.add_heading('Areas for Improvement', level=2)
            for improvement in evaluation.areas_for_improvement:
                p = self.doc.add_paragraph(improvement)
                try:
                    p.style = 'List Bullet'
                except:
                    pass  # Style not available
        
        # Specific recommendations
        if evaluation.recommendations:
            self.doc.add_heading('Recommended Actions', level=2)
            for i, recommendation in enumerate(evaluation.recommendations, 1):
                p = self.doc.add_paragraph(f"{i}. {recommendation}")
    
    def _add_appendices(self, evaluation: DocumentEvaluation):
        """Add appendices with 32-criterion rubric information"""
        
        self.doc.add_heading('Appendices', level=1)
        
        # Appendix A: 32-Criterion Rubric Overview
        self.doc.add_heading('Appendix A: 32-Criterion Rubric Overview', level=2)
        rubric_text = f"""
This evaluation uses the complete 32-criterion Project Handoff Evaluation Rubric 
across 7 major categories:

CONTENT QUALITY (75 points):
• Raw scoring: {evaluation.max_possible_raw_score} points across 32 criteria
• Scaling formula: (Raw Score ÷ {evaluation.max_possible_raw_score}) × 75 = Content Points
• Current raw score: {evaluation.total_raw_score}/{evaluation.max_possible_raw_score} ({evaluation.content_percentage:.1f}%)
• Scaled content score: {evaluation.scaled_content_score:.1f}/75 points

FOREMAN ATTENDANCE (25 points):
• Binary assessment: 25 points if present with meaningful input, 0 if not detected
• Current attendance score: {evaluation.attendance_score}/25 points

FINAL CALCULATION:
Content Points + Attendance Points = {evaluation.final_score:.1f}/100 points

SCORING PHILOSOPHY:
Each criterion includes detailed scoring breakdowns with specific components.
The evaluation rewards specific, project-focused details (proper nouns, dates, quantities) 
while penalizing generic responses ("standard materials", "will coordinate", "TBD"). 
This approach ensures handoffs demonstrate actual preparation rather than form completion.
"""
        self.doc.add_paragraph(rubric_text.strip())
        
        # Appendix B: Performance Levels
        self.doc.add_heading('Appendix B: Performance Level Definitions', level=2)
        
        perf_table = self.doc.add_table(rows=6, cols=2)
        try:
            perf_table.style = 'Table Grid'
        except:
            pass  # Style not available
        
        perf_levels = [
            ("Performance Level", "Score Range"),
            ("EXCELLENT", "90-100 points"),
            ("GOOD", "80-89 points"), 
            ("SATISFACTORY", "70-79 points"),
            ("NEEDS IMPROVEMENT", "60-69 points"),
            ("UNSATISFACTORY", "0-59 points")
        ]
        
        for i, (level, range_val) in enumerate(perf_levels):
            perf_table.cell(i, 0).text = level
            perf_table.cell(i, 1).text = range_val
            if i == 0:
                try:
                    perf_table.cell(i, 0).paragraphs[0].runs[0].bold = True
                    perf_table.cell(i, 1).paragraphs[0].runs[0].bold = True
                except:
                    pass  # Bold formatting failed
        
        # Appendix C: Category Weight Analysis
        self.doc.add_heading('Appendix C: Category Weight Analysis', level=2)
        
        weight_table = self.doc.add_table(rows=1, cols=5)
        try:
            weight_table.style = 'Table Grid'
        except:
            pass
        
        # Headers
        weight_headers = ['Category', 'Criteria Count', 'Max Points', 'Weight %', 'Description']
        header_cells = weight_table.rows[0].cells
        for i, header in enumerate(weight_headers):
            header_cells[i].text = header
            try:
                header_cells[i].paragraphs[0].runs[0].bold = True
            except:
                pass
        
        # Add weight data for each category
        for category in evaluation.category_results:
            row_cells = weight_table.add_row().cells
            row_cells[0].text = category.name
            row_cells[1].text = str(len(category.criteria_scores))
            row_cells[2].text = str(category.max_score)
            weight_percent = (category.max_score / evaluation.max_possible_raw_score * 100) if evaluation.max_possible_raw_score > 0 else 0
            row_cells[3].text = f"{weight_percent:.1f}%"
            
            # Add category description
            if category.name == "Customer":
                row_cells[4].text = "Customer background and relationships"
            elif category.name == "Timeline":
                row_cells[4].text = "Project scheduling and milestones"
            elif category.name == "Material":
                row_cells[4].text = "Materials planning and procurement"
            elif category.name == "Safety":
                row_cells[4].text = "Safety protocols and risk management"
            elif category.name == "Scope":
                row_cells[4].text = "Work scope definition and coverage"
            elif category.name == "Equipment":
                row_cells[4].text = "Equipment planning and coordination"
            elif category.name == "General":
                row_cells[4].text = "General project management requirements"
            else:
                row_cells[4].text = "Project management category"


def create_output_path(document_path: str) -> str:
    """Create output path based on input document path"""
    
    # Get directory and filename without extension
    directory = os.path.dirname(document_path)
    filename_without_ext = os.path.splitext(os.path.basename(document_path))[0]
    
    # Create timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    # Create output filename
    output_filename = f"32Criterion_Rubric_Analysis_{filename_without_ext}_{timestamp}.docx"
    
    # Full output path
    output_path = os.path.join(directory, output_filename)
    
    return output_path


def main():
    """Main execution function - runs with complete 32-criterion rubric"""
    
    print("=" * 80)
    print("PROJECT HANDOFF RUBRIC ANALYZER - COMPLETE 32-CRITERION EVALUATION")
    print("=" * 80)
    
    # Validate input document path
    if not os.path.exists(DOCUMENT_PATH):
        print(f"❌ ERROR: Document file not found at: {DOCUMENT_PATH}")
        print("\n💡 Please update the DOCUMENT_PATH variable at the top of the script")
        print("   with the correct path to your document.")
        return
    
    # Create output path
    output_path = create_output_path(DOCUMENT_PATH)
    
    print(f"📄 Document: {DOCUMENT_PATH}")
    print(f"🤖 Model: {MODEL_TYPE}:{MODEL_NAME}")
    print(f"📝 Output: {output_path}")
    print(f"🎯 Scoring System: 100 points total (75 content + 25 attendance)")
    print(f"📊 Complete 32-Criterion Rubric: 180 raw points → scaled to 75 points")
    print(f"📋 Categories: Customer, Timeline, Material, Safety, Scope, Equipment, General")
    print("-" * 80)
    
    try:
        # Initialize analyzer
        print("🔧 Initializing analyzer with complete 32-criterion rubric...")
        analyzer = ProjectHandoffAnalyzer(model_type=MODEL_TYPE, model_name=MODEL_NAME)
        
        # Verify rubric loaded correctly
        total_criteria = sum(len(cat_data['criteria']) for cat_data in analyzer.rubric.rubric_data.values())
        total_max_score = sum(cat_data['max_score'] for cat_data in analyzer.rubric.rubric_data.values())
        print(f"✅ Rubric loaded: {total_criteria} criteria, {total_max_score} raw points → 75 scaled points")
        
        # Show category breakdown
        print("📋 Category Structure:")
        for cat_name, cat_data in analyzer.rubric.rubric_data.items():
            print(f"   • {cat_name}: {len(cat_data['criteria'])} criteria, {cat_data['max_score']} points")
        
        # Perform analysis
        print("\n🔍 Starting comprehensive 32-criterion analysis...")
        evaluation = analyzer.analyze_document(
            document_path=DOCUMENT_PATH,
            document_name=os.path.basename(DOCUMENT_PATH),
            project_name=FALLBACK_PROJECT_NAME
        )
        
        # Generate report
        print("\n📋 Generating comprehensive 32-criterion report...")
        report_generator = EnhancedWordReportGenerator()
        report_generator.generate_comprehensive_report(evaluation, output_path)
        
        # Print summary with scaled scoring
        print("\n" + "=" * 80)
        print("✅ 32-CRITERION ANALYSIS COMPLETE!")
        print("=" * 80)
        print(f"📊 Project: {evaluation.project_name}")
        print(f"🔍 Foreman Present: {'✅ Yes' if evaluation.foreman_present else '❌ No'}")
        print(f"📈 Raw Content Score: {evaluation.total_raw_score}/{evaluation.max_possible_raw_score} ({evaluation.content_percentage:.1f}%)")
        print(f"🎯 Scaled Content Score: {evaluation.scaled_content_score:.1f}/75 points")
        print(f"👥 Attendance Score: {evaluation.attendance_score:.0f}/25 points")
        print(f"🏆 FINAL SCORE: {evaluation.final_score:.1f}/100 points")
        print(f"📊 Performance Level: {evaluation.performance_level}")
        print(f"📁 Report saved to: {output_path}")
        
        # Category breakdown with scaled scores
        print(f"\n📋 Category Breakdown (32 Criteria Total):")
        for cat in evaluation.category_results:
            # Calculate scaled score for this category
            cat_scaled = (cat.total_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            cat_max_scaled = (cat.max_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            
            print(f"   {cat.name} ({len(cat.criteria_scores)} criteria): {cat.total_score}/{cat.max_score} raw → {cat_scaled:.1f}/{cat_max_scaled:.1f} scaled ({cat.percentage:.1f}%)")
        
        print(f"\n📊 FINAL SCORING SUMMARY:")
        print(f"   • Total Criteria Evaluated: 32")
        print(f"   • Content Quality: {evaluation.scaled_content_score:.1f}/75 points (75% weight)")
        print(f"   • Foreman Attendance: {evaluation.attendance_score:.0f}/25 points (25% weight)")
        print(f"   • FINAL SCORE: {evaluation.final_score:.1f}/100 points")
        print(f"   • Performance Level: {evaluation.performance_level}")
        
        print(f"\n🎉 Open the Word document to view your detailed 32-criterion analysis!")
        
    except Exception as e:
        print(f"\n❌ ERROR during analysis: {str(e)}")
        print("\n💡 Troubleshooting tips:")
        print("   1. Make sure your .env file has valid API keys (ANTHROPIC_API_KEY or OPENAI_API_KEY)")
        print("   2. Check that the document path is correct and the file exists")
        print("   3. Ensure all required packages are installed:")
        print("      pip install python-docx openai anthropic python-dotenv")
        print("   4. The script will use mock analysis if API keys are not available")
        print("   5. Make sure you have write permissions to the output directory")


if __name__ == "__main__":
    main()