#!/usr/bin/env python3
"""
Fixed Project Handoff Rubric Analyzer with LLM Integration - SCALED TO 100 POINTS
Analyzes project handoff documents against the complete 32-criterion rubric
using AI models and generates detailed Word document reports.

SCORING SYSTEM:
- Content Quality: 75 points (75% weight) - scaled from 180 raw points
- Foreman Attendance: 25 points (25% weight)
- TOTAL: 100 points

Requirements:
pip install python-docx openai anthropic python-dotenv

Usage:
Simply modify the DOCUMENT_PATH variable below and run the script directly in VSCode
"""

import os
import json
from typing import Dict, List, Tuple, Any, Optional
from dataclasses import dataclass, field
from datetime import datetime
import re

# Third-party imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

# LLM clients (optional - install as needed)
try:
    import openai
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False

try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

# Load environment variables
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# =============================================================================
# CONFIGURATION - MODIFY THESE SETTINGS
# =============================================================================

# Path to your document (CHANGE THIS to your document path)
DOCUMENT_PATH = "C:\\Users\\kayla.dipaolo\\source\\repos\\provo river water treatment\\rft_analysis_Provo River Water Treatment Turnover.docx"

# Project name (will be auto-extracted from document, or use this as fallback)
FALLBACK_PROJECT_NAME = "Project Analysis"

# Model settings
MODEL_TYPE = "anthropic"  # Options: "openai" or "anthropic"
MODEL_NAME = "claude-3-5-sonnet-20241022"   # Latest model

# =============================================================================

@dataclass
class CriterionScore:
    """Individual criterion scoring breakdown"""
    name: str
    points: int
    max_points: int
    description: str
    justification: str
    supporting_evidence: str = ""

@dataclass
class CategoryResult:
    """Category-level results"""
    name: str
    total_score: int
    max_score: int
    criteria_scores: List[CriterionScore]
    
    @property
    def percentage(self) -> float:
        return (self.total_score / self.max_score * 100) if self.max_score > 0 else 0
    
    @property
    def scaled_score(self) -> float:
        """Scaled score as part of 75-point content system"""
        return (self.total_score / self.max_score * 75) if self.max_score > 0 else 0

@dataclass
class DocumentEvaluation:
    """Complete document evaluation results - SCALED TO 100 POINTS"""
    document_name: str
    project_name: str
    evaluation_date: str
    model_used: str
    category_results: List[CategoryResult] = field(default_factory=list)
    overall_strengths: List[str] = field(default_factory=list)
    areas_for_improvement: List[str] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)
    
    # Foreman attendance tracking
    foreman_present: bool = False
    foreman_names_found: List[str] = field(default_factory=list)
    
    @property
    def total_raw_score(self) -> int:
        """Raw total score from rubric (out of 180)"""
        return sum(cat.total_score for cat in self.category_results)
    
    @property
    def max_possible_raw_score(self) -> int:
        """Maximum raw score possible (180)"""
        return sum(cat.max_score for cat in self.category_results)
    
    @property
    def content_percentage(self) -> float:
        """Raw content score percentage from rubric"""
        return (self.total_raw_score / self.max_possible_raw_score * 100) if self.max_possible_raw_score > 0 else 0
    
    @property
    def scaled_content_score(self) -> float:
        """Scaled content score out of 75 points"""
        return (self.total_raw_score / self.max_possible_raw_score * 75) if self.max_possible_raw_score > 0 else 0
    
    @property
    def attendance_score(self) -> float:
        """Attendance score out of 25 points"""
        return 25.0 if self.foreman_present else 0.0
    
    @property
    def final_score(self) -> float:
        """Final scaled score out of 100 points"""
        return self.scaled_content_score + self.attendance_score
    
    @property
    def weighted_percentage(self) -> float:
        """Weighted percentage (same as final score since it's out of 100)"""
        return self.final_score
    
    @property
    def percentage(self) -> float:
        """Default to final score for backward compatibility"""
        return self.final_score
    
    @property
    def performance_level(self) -> str:
        score = self.final_score
        if score >= 60:
            return "EXCELLENT"
        elif score >= 40:
            return "GOOD"
        elif score >= 20:
            return "SATISFACTORY"
        elif score >= 10:
            return "NEEDS IMPROVEMENT"
        else:
            return "UNSATISFACTORY"

class RubricDefinition:
    """Complete 8-section rubric definition - RUBRIC 2.0"""
    
    def __init__(self):
        self.rubric_data = self._load_complete_rubric()
    
    def _load_complete_rubric(self) -> Dict[str, Any]:
        """Load the complete rubric structure - RUBRIC 2.0 (8 sections)"""
        return {
            "Customer Review": {
                "max_score": 10,
                "criteria": [
                    {
                        "name": "Customer Background Details",
                        "max_points": 10,
                        "description": "Customer organization and contacts named, relationship history documented, specific concerns identified",
                        "scoring_breakdown": {
                            "Customer Identity": {"points": 3, "guide": "3=Specific organization/contacts named | 2=Some names | 1=Generic identification | 0=Not identified"},
                            "Relationship History": {"points": 3, "guide": "3=Specific past projects documented | 2=Some history | 1=Basic mention | 0=No history"},
                            "Current Concerns": {"points": 3, "guide": "3=Specific requirements/concerns identified | 2=Some concerns | 1=Generic concerns | 0=No concerns"},
                            "Evidence Quality": {"points": 1, "guide": "1=Concrete details with proper nouns | 0=Generic statements"}
                        }
                    }
                ]
            },
            "Schedule Review": {
                "max_score": 15,
                "criteria": [
                    {
                        "name": "Project Timeline and Coordination",
                        "max_points": 15,
                        "description": "Specific start/end dates, crew planning, integration with other trades, milestone identification",
                        "scoring_breakdown": {
                            "Timeline Specificity": {"points": 4, "guide": "4=Specific start/end dates | 3=Clear dates | 2=General timeline | 1=Vague dates | 0=TBD"},
                            "Crew Planning": {"points": 4, "guide": "4=Crew size with man-day calculations | 3=Crew details | 2=Basic planning | 1=Minimal details | 0=No planning"},
                            "Trade Integration": {"points": 4, "guide": "4=Detailed coordination with other trades | 3=Good coordination | 2=Basic integration | 1=Limited coordination | 0=No coordination"},
                            "Milestone Management": {"points": 3, "guide": "3=Milestones with responsible parties | 2=Some milestones | 1=Basic milestones | 0=No milestones"}
                        }
                    }
                ]
            },
            "Material Review": {
                "max_score": 15,
                "criteria": [
                    {
                        "name": "Material Planning and Logistics",
                        "max_points": 15,
                        "description": "Specific materials with quantities, suppliers identified, lead times, installation sequence",
                        "scoring_breakdown": {
                            "Material Specificity": {"points": 4, "guide": "4=Specific materials with quantities | 3=Good details | 2=Some specifics | 1=Basic list | 0=Generic materials"},
                            "Supplier Information": {"points": 4, "guide": "4=Suppliers identified by name | 3=Some supplier info | 2=Basic supplier data | 1=Limited info | 0=No supplier info"},
                            "Lead Time Planning": {"points": 4, "guide": "4=Specific lead times with dates | 3=Good timing info | 2=Basic timing | 1=Limited timing | 0=No lead times"},
                            "Installation Process": {"points": 3, "guide": "3=Detailed installation sequence | 2=Some process info | 1=Basic process | 0=No installation planning"}
                        }
                    }
                ]
            },
            "Risk Review": {
                "max_score": 15,
                "criteria": [
                    {
                        "name": "Risk Assessment and Mitigation",
                        "max_points": 15,
                        "description": "Project-specific risks identified, detailed mitigation strategies, site-specific hazard analysis",
                        "scoring_breakdown": {
                            "Risk Identification": {"points": 5, "guide": "5=Project-specific risks (not generic) | 4=Good risk ID | 3=Some specific risks | 2=Mix of generic/specific | 1=Mostly generic | 0=Only generic"},
                            "Mitigation Strategies": {"points": 5, "guide": "5=Detailed mitigation for each risk | 4=Good strategies | 3=Some mitigation | 2=Basic strategies | 1=Limited mitigation | 0=No mitigation"},
                            "Site-Specific Analysis": {"points": 3, "guide": "3=Comprehensive site analysis | 2=Good site considerations | 1=Basic site factors | 0=No site-specific analysis"},
                            "Prevention Methods": {"points": 2, "guide": "2=Specific prevention methods described | 1=Some prevention | 0=No prevention methods"}
                        }
                    }
                ]
            },
            "Equipment Review": {
                "max_score": 10,
                "criteria": [
                    {
                        "name": "Equipment Planning and Coordination",
                        "max_points": 10,
                        "description": "Specific equipment types, costs/rates, provider identification, delivery timeline, responsibility assignments",
                        "scoring_breakdown": {
                            "Equipment Specificity": {"points": 3, "guide": "3=Specific equipment types listed | 2=Good equipment details | 1=Basic equipment info | 0=Generic equipment"},
                            "Cost Information": {"points": 2, "guide": "2=Costs or rental rates provided | 1=Some cost info | 0=No cost information"},
                            "Provider Details": {"points": 2, "guide": "2=Equipment provider identified | 1=Some provider info | 0=No provider information"},
                            "Timeline and Responsibility": {"points": 2, "guide": "2=Delivery timeline and clear responsibilities | 1=Some timing/responsibility info | 0=No timeline/responsibility"},
                            "Analysis Quality": {"points": 1, "guide": "1=Thorough analysis vs 'standard equipment' | 0=Generic responses"}
                        }
                    }
                ]
            },
            "Project Scope Review": {
                "max_score": 20,
                "criteria": [
                    {
                        "name": "Drawing Review and Scope Definition",
                        "max_points": 20,
                        "description": "Evidence of actual drawing review, scope clearly defined, estimating handoff complete, technical documentation",
                        "scoring_breakdown": {
                            "Drawing Review Evidence": {"points": 6, "guide": "6=Specific drawing numbers/dates referenced | 5=Good drawing evidence | 4=Some drawing specifics | 3=Basic drawing review | 2=Limited evidence | 1=Minimal evidence | 0=No evidence"},
                            "Scope Definition": {"points": 5, "guide": "5=Clear scope with inclusions/exclusions | 4=Well-defined scope | 3=Generally clear scope | 2=Some scope clarity | 1=Basic scope | 0=Unclear scope"},
                            "Estimating Handoff": {"points": 5, "guide": "5=Complete estimating handoff documented | 4=Good handoff | 3=Adequate handoff | 2=Some handoff | 1=Basic handoff | 0=No handoff"},
                            "Technical Documentation": {"points": 3, "guide": "3=Technical docs provided (Bluebeam, markups) | 2=Some documentation | 1=Basic docs | 0=No documentation"},
                            "Response Quality": {"points": 1, "guide": "1=Project-specific details vs templates | 0=Templated responses"}
                        }
                    }
                ]
            },
            "Foreman Comments": {
                "max_score": 10,
                "criteria": [
                    {
                        "name": "Foreman Input and Participation",
                        "max_points": 10,
                        "description": "Foreman present and active, specific suggestions/concerns, actionable items, expertise applied",
                        "scoring_breakdown": {
                            "Participation Level": {"points": 4, "guide": "4=Active participation with expertise | 3=Good participation | 2=Some participation | 1=Limited participation | 0=No participation"},
                            "Specific Input": {"points": 3, "guide": "3=Specific suggestions/concerns documented | 2=Some specific input | 1=Basic input | 0=No specific input"},
                            "Actionable Items": {"points": 2, "guide": "2=Clear actionable items identified | 1=Some actionable items | 0=No actionable items"},
                            "Evidence of Expertise": {"points": 1, "guide": "1=Clear evidence of foreman expertise applied | 0=No evidence of expertise"}
                        }
                    }
                ]
            },
            "Payroll Review": {
                "max_score": 5,
                "criteria": [
                    {
                        "name": "Payroll Setup and Verification",
                        "max_points": 5,
                        "description": "Specific payroll attributes confirmed, labor codes verified, setup completion documented",
                        "scoring_breakdown": {
                            "Payroll Attributes": {"points": 2, "guide": "2=Specific payroll attributes confirmed | 1=Some attributes | 0=Generic review"},
                            "Labor Codes": {"points": 2, "guide": "2=Labor codes and classifications verified | 1=Some verification | 0=No verification"},
                            "Setup Completion": {"points": 1, "guide": "1=Setup completion documented | 0=Not confirmed"}
                        }
                    }
                ]
            }
        }

class LLMAnalyzer:
    """Handles LLM-based document analysis"""
    
    def __init__(self, model_type: str = "anthropic", model_name: str = "claude-3-5-sonnet-20241022"):
        self.model_type = model_type
        self.model_name = model_name
        self._setup_client()
    
    def _setup_client(self):
        """Setup the appropriate LLM client"""
        if self.model_type == "openai" and HAS_OPENAI:
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                print("Warning: OPENAI_API_KEY not found in environment variables.")
                self.client = None
            else:
                self.client = openai.OpenAI(api_key=api_key)
        elif self.model_type == "anthropic" and HAS_ANTHROPIC:
            api_key = os.getenv("ANTHROPIC_API_KEY")
            if not api_key:
                print("Warning: ANTHROPIC_API_KEY not found in environment variables.")
                self.client = None
            else:
                self.client = anthropic.Anthropic(api_key=api_key)
        else:
            self.client = None
            print(f"Warning: {self.model_type} client not available. Using mock analysis.")
    
    def analyze_criterion(self, document_content: str, criterion_data: Dict[str, Any]) -> CriterionScore:
        """Analyze a single criterion using LLM with retry logic"""
        
        if self.client is None:
            # Mock analysis for demonstration
            return self._mock_analysis(criterion_data)
        
        prompt = self._build_analysis_prompt(criterion_data, document_content)
        
        # Retry logic for API overload errors
        max_retries = 3
        base_delay = 2  # seconds
        
        for attempt in range(max_retries):
            try:
                if self.model_type == "openai":
                    response = self.client.chat.completions.create(
                        model=self.model_name,
                        messages=[
                            {"role": "system", "content": "You are an expert project management analyst evaluating handoff documents against specific criteria."},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.1
                    )
                    result_text = response.choices[0].message.content
                    
                elif self.model_type == "anthropic":
                    response = self.client.messages.create(
                        model=self.model_name,
                        max_tokens=1000,
                        temperature=0.1,
                        messages=[
                            {"role": "user", "content": prompt}
                        ]
                    )
                    
                    if hasattr(response, 'content') and response.content and len(response.content) > 0:
                        if hasattr(response.content[0], 'text'):
                            result_text = response.content[0].text
                        else:
                            result_text = str(response.content[0])
                    else:
                        print(f"Warning: Unexpected response format from Anthropic API")
                        return self._mock_analysis(criterion_data)
                
                return self._parse_llm_response(result_text, criterion_data)
                
            except Exception as e:
                error_str = str(e)
                
                # Check for overload error (529)
                if "529" in error_str or "overloaded" in error_str.lower():
                    if attempt < max_retries - 1:  # Don't sleep on last attempt
                        delay = base_delay * (2 ** attempt)  # Exponential backoff
                        print(f"   API overloaded, retrying in {delay} seconds... (attempt {attempt + 1}/{max_retries})")
                        import time
                        time.sleep(delay)
                        continue
                    else:
                        print(f"   API overloaded after {max_retries} attempts, using mock analysis")
                        return self._mock_analysis(criterion_data)
                else:
                    # For other errors, fail immediately
                    print(f"Error in LLM analysis for '{criterion_data['name']}': {e}")
                    return self._mock_analysis(criterion_data)
        
        # Fallback if all retries failed
        return self._mock_analysis(criterion_data)
    
    def _build_analysis_prompt(self, criterion_data: Dict[str, Any], document_content: str) -> str:
        """Build the analysis prompt for the LLM"""
        
        prompt = f"""
Analyze the following project handoff document against this specific section:

SECTION: {criterion_data['name']}
DESCRIPTION: {criterion_data['description']}
MAX POINTS: {criterion_data['max_points']}

SCORING BREAKDOWN:
"""
        
        for component, details in criterion_data['scoring_breakdown'].items():
            prompt += f"- {component} ({details['points']} pts): {details['guide']}\n"
        
        prompt += f"""

DOCUMENT TO ANALYZE:
{document_content[:3000]}...

SCORING FOCUS:
- Reward SPECIFIC details: proper nouns, exact dates, quantities, dollar amounts
- Penalize GENERIC responses: "standard materials", "typical process", "will coordinate"
- Look for PROJECT-SPECIFIC information rather than templated language
- Evidence of actual preparation vs. form completion

Please provide your analysis in this exact format:
TOTAL_SCORE: [number out of {criterion_data['max_points']}]
JUSTIFICATION: [detailed explanation focusing on specificity vs. generic content]
SUPPORTING_EVIDENCE: [specific quotes or references from the document]

Focus on rewarding concrete, project-specific details and penalizing vague, generic responses.
"""
        
        return prompt
    
    def _parse_llm_response(self, response_text: str, criterion_data: Dict[str, Any]) -> CriterionScore:
        """Parse the LLM response into a structured result"""
        
        # Extract score
        score_match = re.search(r'TOTAL_SCORE:\s*(\d+)', response_text)
        score = int(score_match.group(1)) if score_match else 0
        
        # Extract justification
        just_match = re.search(r'JUSTIFICATION:\s*(.*?)(?=SUPPORTING_EVIDENCE:|$)', response_text, re.DOTALL)
        justification = just_match.group(1).strip() if just_match else "No justification provided"
        
        # Extract supporting evidence
        evidence_match = re.search(r'SUPPORTING_EVIDENCE:\s*(.*?)$', response_text, re.DOTALL)
        evidence = evidence_match.group(1).strip() if evidence_match else ""
        
        return CriterionScore(
            name=criterion_data['name'],
            points=min(score, criterion_data['max_points']),  # Ensure not over max
            max_points=criterion_data['max_points'],
            description=criterion_data['description'],
            justification=justification,
            supporting_evidence=evidence
        )
    
    def _mock_analysis(self, criterion_data: Dict[str, Any]) -> CriterionScore:
        """Provide mock analysis when LLM is not available"""
        
        # Simple mock scoring (for demonstration)
        mock_score = max(1, int(criterion_data['max_points'] * 0.7))  # 70% score, minimum 1
        
        return CriterionScore(
            name=criterion_data['name'],
            points=mock_score,
            max_points=criterion_data['max_points'],
            description=criterion_data['description'],
            justification="Mock analysis - LLM client not available or API key missing",
            supporting_evidence="Mock evidence for demonstration purposes"
        )

class ProjectHandoffAnalyzer:
    """Main analyzer class that coordinates the analysis"""
    
    def __init__(self, model_type: str = "anthropic", model_name: str = "claude-3-5-sonnet-20241022"):
        self.rubric = RubricDefinition()
        self.llm_analyzer = LLMAnalyzer(model_type, model_name)
    
    def _detect_foreman_attendance(self, document_content: str) -> Tuple[bool, List[str]]:
        """Use AI to detect if foreman attended the meeting based on transcript analysis"""
        
        if self.llm_analyzer.client is None:
            # Fallback for mock analysis
            return True, ["Mock: Foreman attendance assumed for demonstration"]
        
        prompt = f"""
Analyze this project handoff meeting transcript to determine if a FOREMAN was present and participating.

A foreman is typically:
- The field supervisor or construction lead
- Someone who provides installation insights, field experience, or practical input
- Someone who discusses field conditions, equipment needs, or installation challenges
- Someone who gives feedback on constructability or field execution

Look for evidence such as:
- Someone providing field/installation expertise or suggestions
- Discussion of field conditions, access, equipment needs from someone with hands-on experience
- Practical input about installation challenges, safety, or site logistics
- Someone asking/answering questions about actual construction execution
- References to someone who will be managing the field work

TRANSCRIPT TO ANALYZE:
{document_content[:2000]}...

Provide your assessment in this exact format:
FOREMAN_PRESENT: [YES or NO]
CONFIDENCE: [LOW, MEDIUM, HIGH]
EVIDENCE: [specific quotes or observations that support your determination]
REASONING: [explain why you believe a foreman was or wasn't present]

Focus on the ROLE and CONTRIBUTIONS rather than job titles or names.
"""
        
        try:
            if self.llm_analyzer.model_type == "openai":
                response = self.llm_analyzer.client.chat.completions.create(
                    model=self.llm_analyzer.model_name,
                    messages=[
                        {"role": "system", "content": "You are an expert construction project analyst who can identify team roles and participation based on meeting content."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1
                )
                result_text = response.choices[0].message.content

            elif self.llm_analyzer.model_type == "anthropic":
                response = self.llm_analyzer.client.messages.create(
                    model=self.llm_analyzer.model_name,
                    max_tokens=800,
                    temperature=0.1,
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )
                
                if hasattr(response, 'content') and response.content and len(response.content) > 0:
                    if hasattr(response.content[0], 'text'):
                        result_text = response.content[0].text
                    else:
                        result_text = str(response.content[0])
                else:
                    print(f"Warning: Unexpected response format from Anthropic API")
                    return False, ["API response format error"]

            # Parse the LLM response
            foreman_present = "YES" in result_text.upper() and "FOREMAN_PRESENT:" in result_text.upper()

            # Extract evidence and reasoning
            evidence_match = re.search(r'EVIDENCE:\s*(.*?)(?=REASONING:|$)', result_text, re.DOTALL)
            evidence = evidence_match.group(1).strip() if evidence_match else "No evidence extracted"

            reasoning_match = re.search(r'REASONING:\s*(.*?)$', result_text, re.DOTALL)
            reasoning = reasoning_match.group(1).strip() if reasoning_match else "No reasoning provided"

            confidence_match = re.search(r'CONFIDENCE:\s*(\w+)', result_text, re.IGNORECASE)
            confidence = confidence_match.group(1) if confidence_match else "UNKNOWN"

            analysis_details = [
                f"Confidence: {confidence}",
                f"Evidence: {evidence[:100]}..." if len(evidence) > 100 else f"Evidence: {evidence}",
                f"Reasoning: {reasoning[:100]}..." if len(reasoning) > 100 else f"Reasoning: {reasoning}"
            ]

            return foreman_present, analysis_details

        except Exception as e:
            print(f"Error in foreman detection: {e}")
            return False, [f"Error occurred during analysis: {str(e)}"]
    
    def analyze_document(self, document_path: str, document_name: str = None, 
                        project_name: str = "Project Analysis") -> DocumentEvaluation:
        """Analyze a complete document against the rubric"""
        
        # Read document based on file type
        document_content = self._read_document(document_path)
        
        if document_name is None:
            document_name = os.path.basename(document_path)
        
        # Auto-extract project name
        extracted_project_name = self._extract_project_name(document_content)
        if extracted_project_name != "Unknown Project":
            project_name = extracted_project_name
        
        # Initialize evaluation
        evaluation = DocumentEvaluation(
            document_name=document_name,
            project_name=project_name,
            evaluation_date=datetime.now().strftime("%Y-%m-%d"),
            model_used=f"{self.llm_analyzer.model_type}:{self.llm_analyzer.model_name}"
        )
        
        # Detect foreman attendance
        foreman_present, foreman_indicators = self._detect_foreman_attendance(document_content)
        evaluation.foreman_present = foreman_present
        evaluation.foreman_names_found = foreman_indicators
        
        print(f"🔍 Foreman Detection: {'✅ Present' if foreman_present else '❌ Not detected'}")
        if foreman_indicators:
            print(f"   Analysis: {foreman_indicators[0] if foreman_indicators else 'No details'}")
        
        # Analyze each category
        total_sections = 0
        for category_name, category_data in self.rubric.rubric_data.items():
            category_result = self._analyze_category(
                document_content, category_name, category_data
            )
            evaluation.category_results.append(category_result)
            total_sections += len(category_data['criteria'])
        
        print(f"📊 Analyzed {total_sections} sections across {len(evaluation.category_results)} categories")
        
        # Generate overall insights
        evaluation.overall_strengths = self._identify_strengths(evaluation)
        evaluation.areas_for_improvement = self._identify_improvements(evaluation)
        evaluation.recommendations = self._generate_recommendations(evaluation)
        
        return evaluation
    
    def _read_document(self, document_path: str) -> str:
        """Read document content based on file type"""
        
        try:
            file_ext = document_path.lower().split('.')[-1]
            
            if file_ext == 'docx':
                # Read Word document
                doc = Document(document_path)
                document_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
            elif file_ext == 'json':
                # Read JSON file and convert to readable text
                with open(document_path, 'r', encoding='utf-8') as f:
                    json_data = json.load(f)
                document_content = self._convert_json_to_text(json_data)
                
            elif file_ext == 'csv':
                # Read CSV file
                import csv
                document_content = ""
                with open(document_path, 'r', encoding='utf-8') as f:
                    csv_reader = csv.reader(f)
                    for row in csv_reader:
                        document_content += " | ".join(row) + "\n"
                        
            else:
                # Read as text file
                with open(document_path, 'r', encoding='utf-8') as f:
                    document_content = f.read()
            
            return document_content
            
        except Exception as e:
            raise Exception(f"Error reading document '{document_path}': {str(e)}")
    
    def _convert_json_to_text(self, json_data: Any, indent_level: int = 0) -> str:
        """Convert JSON data to readable text format for analysis"""
        
        indent = "  " * indent_level
        text_parts = []
        
        if isinstance(json_data, dict):
            for key, value in json_data.items():
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{indent}{key}:")
                    text_parts.append(self._convert_json_to_text(value, indent_level + 1))
                else:
                    text_parts.append(f"{indent}{key}: {value}")
        
        elif isinstance(json_data, list):
            for i, item in enumerate(json_data):
                if isinstance(item, (dict, list)):
                    text_parts.append(f"{indent}Item {i + 1}:")
                    text_parts.append(self._convert_json_to_text(item, indent_level + 1))
                else:
                    text_parts.append(f"{indent}- {item}")
        
        else:
            text_parts.append(f"{indent}{json_data}")
        
        return "\n".join(text_parts)

    def _analyze_category(self, document_content: str, category_name: str, 
                         category_data: Dict[str, Any]) -> CategoryResult:
        """Analyze all criteria in a category"""
        
        criteria_scores = []
        
        for criterion_data in category_data['criteria']:
            print(f"  Analyzing: {criterion_data['name']}")
            score = self.llm_analyzer.analyze_criterion(document_content, criterion_data)
            criteria_scores.append(score)
        
        total_score = sum(score.points for score in criteria_scores)
        max_score = sum(score.max_points for score in criteria_scores)
        
        return CategoryResult(
            name=category_name,
            total_score=total_score,
            max_score=max_score,
            criteria_scores=criteria_scores
        )
    
    def _extract_project_name(self, document_content: str) -> str:
        """Extract project name from document content using various patterns"""
        
        # Patterns to look for (in order of preference)
        patterns = [
            # Direct project name patterns
            r'project\s*name[:\-\s]+(.+?)(?:\n|$)',
            r'project[:\-\s]+(.+?)(?:\n|$)',
            r'job\s*name[:\-\s]+(.+?)(?:\n|$)',
            r'job[:\-\s]+(.+?)(?:\n|$)',
            r'site[:\-\s]+(.+?)(?:\n|$)',
            r'client[:\-\s]+(.+?)(?:\n|$)',
            r'customer[:\-\s]+(.+?)(?:\n|$)',
        ]
        
        # Try each pattern
        for pattern in patterns:
            matches = re.findall(pattern, document_content, re.IGNORECASE | re.MULTILINE)
            if matches:
                # Clean up the match
                project_name = matches[0].strip()
                # Remove common prefixes/suffixes
                project_name = re.sub(r'^(the|a|an)\s+', '', project_name, flags=re.IGNORECASE)
                project_name = re.sub(r'\s+(project|handoff|document|summary)$', '', project_name, flags=re.IGNORECASE)
                
                if project_name and len(project_name) > 2:  # Make sure we have a valid project name
                    return project_name
        
        return "Unknown Project"
    
    def _identify_strengths(self, evaluation: DocumentEvaluation) -> List[str]:
        """Identify overall strengths from the analysis"""
        strengths = []
        
        # Find high-performing categories
        high_performers = [cat for cat in evaluation.category_results if cat.percentage >= 80]
        
        if high_performers:
            strengths.append(f"Strong performance in {', '.join([cat.name for cat in high_performers])}")
        
        # Find perfect scores
        perfect_scores = []
        for cat in evaluation.category_results:
            for crit in cat.criteria_scores:
                if crit.points == crit.max_points:
                    perfect_scores.append(crit.name)
        
        if perfect_scores:
            strengths.append(f"Perfect scores achieved in: {', '.join(perfect_scores[:3])}...")
        
        if not strengths:
            strengths.append("Analysis completed successfully with comprehensive evaluation")
        
        return strengths
    
    def _identify_improvements(self, evaluation: DocumentEvaluation) -> List[str]:
        """Identify areas needing improvement"""
        improvements = []
        
        # Find low-performing categories
        low_performers = [cat for cat in evaluation.category_results if cat.percentage < 60]
        
        if low_performers:
            improvements.append(f"Significant improvement needed in {', '.join([cat.name for cat in low_performers])}")
        
        # Find zero scores
        zero_scores = []
        for cat in evaluation.category_results:
            for crit in cat.criteria_scores:
                if crit.points == 0:
                    zero_scores.append(crit.name)
        
        if zero_scores:
            improvements.append(f"Critical gaps identified in: {', '.join(zero_scores[:3])}...")
        
        if not improvements:
            improvements.append("Continue maintaining current quality standards")
        
        return improvements
    
    def _generate_recommendations(self, evaluation: DocumentEvaluation) -> List[str]:
        """Generate specific recommendations"""
        recommendations = [
            "Develop comprehensive implementation procedures for all major categories",
            "Create detailed checklists for each criterion to ensure complete coverage",
            "Establish regular review cycles to maintain document quality",
            "Implement cross-functional review process for complex projects"
        ]
        
        # Add specific recommendations based on performance
        if evaluation.final_score < 70:
            recommendations.insert(0, "Immediate comprehensive review and enhancement required across all categories")
        
        return recommendations

class EnhancedWordReportGenerator:
    """Enhanced Word document generator with professional formatting - SCALED SCORING"""
    
    def __init__(self):
        self.doc = None
    
    def generate_comprehensive_report(self, evaluation: DocumentEvaluation, output_path: str):
        """Generate a comprehensive, professionally formatted report"""
        
        self.doc = Document()
        self._setup_styles()
        
        # Generate all sections
        self._add_title_page(evaluation)
        self._add_executive_summary(evaluation)
        self._add_category_analysis(evaluation)
        self._add_detailed_criteria_analysis(evaluation)
        self._add_recommendations_section(evaluation)
        self._add_appendices(evaluation)
        
        # Save document
        self.doc.save(output_path)
        print(f"📄 Comprehensive report saved to: {output_path}")
    
    def _setup_styles(self):
        """Setup professional document styles"""
        styles = self.doc.styles
        
        # Title style
        try:
            if 'Report Title' not in [s.name for s in styles]:
                title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.name = 'Calibri'
                title_style.font.size = Pt(24)
                title_style.font.bold = True
                title_style.font.color.rgb = RGBColor(0, 51, 102)
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass  # Style already exists or can't be created
        
        # Subtitle style
        try:
            if 'Report Subtitle' not in [s.name for s in styles]:
                subtitle_style = styles.add_style('Report Subtitle', WD_STYLE_TYPE.PARAGRAPH)
                subtitle_style.font.name = 'Calibri'
                subtitle_style.font.size = Pt(14)
                subtitle_style.font.color.rgb = RGBColor(89, 89, 89)
                subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass  # Style already exists or can't be created
    
    def _add_title_page(self, evaluation: DocumentEvaluation):
        """Add professional title page with scaled scoring"""
        
        # Title
        try:
            title = self.doc.add_paragraph("Project Handoff Evaluation Report", style='Report Title')
        except:
            title = self.doc.add_paragraph("Project Handoff Evaluation Report")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Subtitle
        try:
            subtitle = self.doc.add_paragraph(evaluation.document_name, style='Report Subtitle')
        except:
            subtitle = self.doc.add_paragraph(evaluation.document_name)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Key metrics table with scaled scoring
        table = self.doc.add_table(rows=10, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        metrics = [
            ("Project", evaluation.project_name),
            ("Evaluation Date", evaluation.evaluation_date),
            ("Model Used", evaluation.model_used),
            ("Foreman Present", "✅ Yes" if evaluation.foreman_present else "❌ No"),
            ("Raw Content Score", f"{evaluation.total_raw_score}/{evaluation.max_possible_raw_score} ({evaluation.content_percentage:.1f}%)"),
            ("Scaled Content Score", f"{evaluation.scaled_content_score:.1f}/75 points"),
            ("Attendance Score", f"{evaluation.attendance_score:.0f}/25 points"),
            ("Final Score", f"{evaluation.final_score:.1f}/100 points"),
            ("Performance Level", evaluation.performance_level)
        ]
        
        for i, (label, value) in enumerate(metrics):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            try:
                table.cell(i, 0).paragraphs[0].runs[0].bold = True
            except:
                pass  # Bold formatting failed
        
        self.doc.add_page_break()
    
    def _add_executive_summary(self, evaluation: DocumentEvaluation):
        """Add executive summary with scaled scoring explanation"""
        
        self.doc.add_heading('Executive Summary', level=1)
        
        # Performance overview with scaled scoring
        performance_para = self.doc.add_paragraph()
        performance_para.add_run(f"Overall Performance: ").bold = True
        performance_para.add_run(
            f"This document achieved a final score of {evaluation.final_score:.1f}/100 points "
            f"(Content: {evaluation.scaled_content_score:.1f}/75 + Attendance: {evaluation.attendance_score:.0f}/25), "
            f"earning a "
        )
        perf_run = performance_para.add_run(evaluation.performance_level)
        perf_run.bold = True
        try:
            if evaluation.performance_level in ['EXCELLENT', 'GOOD']:
                perf_run.font.color.rgb = RGBColor(0, 128, 0)
            elif evaluation.performance_level == 'SATISFACTORY':
                perf_run.font.color.rgb = RGBColor(255, 165, 0)
            else:
                perf_run.font.color.rgb = RGBColor(255, 0, 0)
        except:
            pass  # Color formatting failed
        performance_para.add_run(" rating.")
        
        # Scaling explanation
        self.doc.add_heading('Scoring System Overview', level=2)
        scaling_para = self.doc.add_paragraph()
        scaling_para.add_run("Scoring Methodology: ").bold = True
        scaling_para.add_run(
            f"The evaluation uses a scaled 100-point system. Raw content scores ({evaluation.total_raw_score}/{evaluation.max_possible_raw_score}) "
            f"are scaled to 75 points, representing 75% of the total score. Foreman attendance contributes 25 points (25% weight). "
            f"This ensures consistent, interpretable results across all evaluations."
        )
        
        # Foreman attendance summary
        self.doc.add_heading('Foreman Attendance Analysis', level=2)
        foreman_para = self.doc.add_paragraph()
        if evaluation.foreman_present:
            foreman_para.add_run("✅ Foreman Present: ").bold = True
            foreman_para.add_run(f"Foreman attendance detected. Analysis details: {evaluation.foreman_names_found[0] if evaluation.foreman_names_found else 'No details available'}. ")
            foreman_para.add_run("Earned full 25/25 points for attendance.")
        else:
            foreman_para.add_run("❌ Foreman Not Detected: ").bold = True
            foreman_para.add_run("No clear foreman attendance indicators found in the transcript. ")
            foreman_para.add_run("0/25 points awarded for attendance. Consider ensuring foreman participation is documented.")
        
        # Category performance summary with scaled scores
        self.doc.add_heading('Content Quality Analysis', level=2)
        
        table = self.doc.add_table(rows=1, cols=5)
        try:
            table.style = 'Table Grid'
        except:
            pass  # Style not available
        
        # Headers
        headers = ['Category', 'Raw Score', 'Raw Max', 'Scaled Score', 'Percentage']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            try:
                header_cells[i].paragraphs[0].runs[0].bold = True
            except:
                pass
        
        # Add data rows with scaled scores
        for category in evaluation.category_results:
            row_cells = table.add_row().cells
            row_cells[0].text = category.name
            row_cells[1].text = str(category.total_score)
            row_cells[2].text = str(category.max_score)
            
            # Calculate scaled score for this category
            scaled_score = (category.total_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            row_cells[3].text = f"{scaled_score:.1f}"
            row_cells[4].text = f"{category.percentage:.1f}%"
    
    def _add_category_analysis(self, evaluation: DocumentEvaluation):
        """Add detailed category-by-category analysis with scaled scoring"""
        
        self.doc.add_heading('Category Analysis', level=1)
        
        for category in evaluation.category_results:
            scaled_score = (category.total_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            
            self.doc.add_heading(
                f'{category.name} ({category.total_score}/{category.max_score} raw | {scaled_score:.1f}/75 scaled - {category.percentage:.1f}%)', 
                level=2
            )
            
            # Category summary
            if category.percentage >= 80:
                summary = "Strong performance with comprehensive coverage."
            elif category.percentage >= 60:
                summary = "Adequate performance with room for improvement."
            else:
                summary = "Significant gaps requiring immediate attention."
            
            self.doc.add_paragraph(summary)
            
            # Individual criteria in this category
            for criterion in category.criteria_scores:
                self.doc.add_heading(f'{criterion.name} ({criterion.points}/{criterion.max_points})', level=3)
                self.doc.add_paragraph(f"Description: {criterion.description}")
                self.doc.add_paragraph(f"Assessment: {criterion.justification}")
                
                if criterion.supporting_evidence:
                    self.doc.add_paragraph(f"Evidence: {criterion.supporting_evidence}")
    
    def _add_detailed_criteria_analysis(self, evaluation: DocumentEvaluation):
        """Add comprehensive criteria analysis"""
        
        self.doc.add_heading('Detailed Criterion Analysis', level=1)
        
        criterion_number = 1
        for category in evaluation.category_results:
            for criterion in category.criteria_scores:
                self.doc.add_heading(f'{criterion_number}. {criterion.name}', level=2)
                
                # Score badge
                score_para = self.doc.add_paragraph()
                score_run = score_para.add_run(f"Score: {criterion.points}/{criterion.max_points}")
                score_run.bold = True
                
                try:
                    if criterion.points == criterion.max_points:
                        score_run.font.color.rgb = RGBColor(0, 128, 0)
                    elif criterion.points >= criterion.max_points * 0.7:
                        score_run.font.color.rgb = RGBColor(255, 165, 0) 
                    else:
                        score_run.font.color.rgb = RGBColor(255, 0, 0)
                except:
                    pass  # Color formatting failed
                
                # Analysis details
                self.doc.add_paragraph(f"Criterion: {criterion.description}")
                self.doc.add_paragraph(f"Analysis: {criterion.justification}")
                
                if criterion.supporting_evidence:
                    self.doc.add_paragraph(f"Supporting Evidence: {criterion.supporting_evidence}")
                
                criterion_number += 1
    
    def _add_recommendations_section(self, evaluation: DocumentEvaluation):
        """Add comprehensive recommendations"""
        
        self.doc.add_heading('Recommendations & Action Items', level=1)
        
        # Strengths
        if evaluation.overall_strengths:
            self.doc.add_heading('Key Strengths', level=2)
            for strength in evaluation.overall_strengths:
                p = self.doc.add_paragraph(strength)
                try:
                    p.style = 'List Bullet'
                except:
                    pass  # Style not available
        
        # Areas for improvement  
        if evaluation.areas_for_improvement:
            self.doc.add_heading('Areas for Improvement', level=2)
            for improvement in evaluation.areas_for_improvement:
                p = self.doc.add_paragraph(improvement)
                try:
                    p.style = 'List Bullet'
                except:
                    pass  # Style not available
        
        # Specific recommendations
        if evaluation.recommendations:
            self.doc.add_heading('Recommended Actions', level=2)
            for i, recommendation in enumerate(evaluation.recommendations, 1):
                p = self.doc.add_paragraph(f"{i}. {recommendation}")
    
    def _add_appendices(self, evaluation: DocumentEvaluation):
        """Add appendices with supporting information including scaling explanation"""
        
        self.doc.add_heading('Appendices', level=1)
        
        # Appendix A: Scoring Guide
        self.doc.add_heading('Appendix A: Scoring Methodology', level=2)
        methodology_text = f"""
This evaluation uses the Project Handoff Evaluation Rubric Version 2.0 
with 8 sections focused on content specificity and quality:

CONTENT QUALITY (75 points):
• Raw scoring: 100 points across 8 sections
• Scaling formula: (Raw Score ÷ 100) × 75 = Content Points
• Current raw score: {evaluation.total_raw_score}/100 ({evaluation.content_percentage:.1f}%)
• Scaled content score: {evaluation.scaled_content_score:.1f}/75 points

FOREMAN ATTENDANCE (25 points):
• Binary assessment: 25 points if present with meaningful input, 0 if not detected
• Current attendance score: {evaluation.attendance_score}/25 points

FINAL CALCULATION:
Content Points + Attendance Points = {evaluation.final_score:.1f}/100 points

SCORING PHILOSOPHY:
The evaluation rewards specific, project-focused details (proper nouns, dates, quantities) 
while penalizing generic responses ("standard materials", "will coordinate", "TBD"). 
This approach ensures handoffs demonstrate actual preparation rather than form completion.
"""
        self.doc.add_paragraph(methodology_text.strip())
        
        # Appendix B: Performance Levels
        self.doc.add_heading('Appendix B: Performance Level Definitions', level=2)
        
        perf_table = self.doc.add_table(rows=6, cols=2)
        try:
            perf_table.style = 'Table Grid'
        except:
            pass  # Style not available
        
        perf_levels = [
            ("Performance Level", "Score Range"),
            ("EXCELLENT", "60-100 points"),
            ("GOOD", "40-59 points"), 
            ("SATISFACTORY", "20-39 points"),
            ("NEEDS IMPROVEMENT", "10-19 points"),
            ("UNSATISFACTORY", "0-9 points")
        ]
        
        for i, (level, range_val) in enumerate(perf_levels):
            perf_table.cell(i, 0).text = level
            perf_table.cell(i, 1).text = range_val
            if i == 0:
                try:
                    perf_table.cell(i, 0).paragraphs[0].runs[0].bold = True
                    perf_table.cell(i, 1).paragraphs[0].runs[0].bold = True
                except:
                    pass  # Bold formatting failed
        
        # Appendix C: Category Scaling Breakdown
        self.doc.add_heading('Appendix C: Category Scaling Breakdown', level=2)
        
        scaling_table = self.doc.add_table(rows=1, cols=4)
        try:
            scaling_table.style = 'Table Grid'
        except:
            pass
        
        # Headers
        scaling_headers = ['Category', 'Raw Points', 'Scaled Points', 'Proportion of Total']
        header_cells = scaling_table.rows[0].cells
        for i, header in enumerate(scaling_headers):
            header_cells[i].text = header
            try:
                header_cells[i].paragraphs[0].runs[0].bold = True
            except:
                pass
        
        # Add scaling data for each category
        for category in evaluation.category_results:
            row_cells = scaling_table.add_row().cells
            row_cells[0].text = category.name
            row_cells[1].text = f"{category.total_score}/{category.max_score}"
            
            # Calculate scaled points for this category
            category_scaled = (category.total_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            category_max_scaled = (category.max_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            
            row_cells[2].text = f"{category_scaled:.1f}/{category_max_scaled:.1f}"
            row_cells[3].text = f"{(category.max_score / evaluation.max_possible_raw_score * 100):.1f}%"


def create_output_path(document_path: str) -> str:
    """Create output path based on input document path"""
    
    # Get directory and filename without extension
    directory = os.path.dirname(document_path)
    filename_without_ext = os.path.splitext(os.path.basename(document_path))[0]
    
    # Create timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    # Create output filename
    output_filename = f"Scaled_Rubric_Analysis_{filename_without_ext}_{timestamp}.docx"
    
    # Full output path
    output_path = os.path.join(directory, output_filename)
    
    return output_path


def main():
    """Main execution function - runs directly with scaled 100-point scoring"""
    
    print("=" * 70)
    print("PROJECT HANDOFF RUBRIC ANALYZER - SCALED TO 100 POINTS")
    print("=" * 70)
    
    # Validate input document path
    if not os.path.exists(DOCUMENT_PATH):
        print(f"❌ ERROR: Document file not found at: {DOCUMENT_PATH}")
        print("\n💡 Please update the DOCUMENT_PATH variable at the top of the script")
        print("   with the correct path to your document.")
        return
    
    # Create output path
    output_path = create_output_path(DOCUMENT_PATH)
    
    print(f"📄 Document: {DOCUMENT_PATH}")
    print(f"🤖 Model: {MODEL_TYPE}:{MODEL_NAME}")
    print(f"📝 Output: {output_path}")
    print(f"🎯 Scoring System: 100 points total (75 content + 25 attendance)")
    print(f"📊 Raw Content Rubric: 180 points (32 criteria) → scaled to 75 points")
    print("-" * 70)
    
    try:
        # Initialize analyzer
        print("🔧 Initializing analyzer with specificity-focused 8-section rubric...")
        analyzer = ProjectHandoffAnalyzer(model_type=MODEL_TYPE, model_name=MODEL_NAME)
        
        # Verify rubric loaded correctly
        total_sections = sum(len(cat_data['criteria']) for cat_data in analyzer.rubric.rubric_data.values())
        total_max_score = sum(cat_data['max_score'] for cat_data in analyzer.rubric.rubric_data.values())
        print(f"✅ Rubric loaded: {total_sections} sections, {total_max_score} raw points → 75 scaled points")
        
        # Perform analysis
        print("🔍 Starting document analysis...")
        evaluation = analyzer.analyze_document(
            document_path=DOCUMENT_PATH,
            document_name=os.path.basename(DOCUMENT_PATH),
            project_name=FALLBACK_PROJECT_NAME
        )
        
        # Generate report
        print("📋 Generating comprehensive specificity-based report...")
        report_generator = EnhancedWordReportGenerator()
        report_generator.generate_comprehensive_report(evaluation, output_path)
        
        # Print summary with scaled scoring
        print("\n" + "=" * 70)
        print("✅ ANALYSIS COMPLETE - SCALED SCORING RESULTS!")
        print("=" * 70)
        print(f"📊 Project: {evaluation.project_name}")
        print(f"🔍 Foreman Present: {'✅ Yes' if evaluation.foreman_present else '❌ No'}")
        print(f"📈 Raw Content Score: {evaluation.total_raw_score}/{evaluation.max_possible_raw_score} ({evaluation.content_percentage:.1f}%)")
        print(f"🎯 Scaled Content Score: {evaluation.scaled_content_score:.1f}/75 points")
        print(f"👥 Attendance Score: {evaluation.attendance_score:.0f}/25 points")
        print(f"🏆 FINAL SCORE: {evaluation.final_score:.1f}/100 points")
        print(f"📊 Performance Level: {evaluation.performance_level}")
        print(f"📁 Report saved to: {output_path}")
        print("\n🎉 Open the Word document to view your detailed scaled analysis!")
        
        # Category breakdown with scaled scores
        print(f"\n📋 Section Breakdown (Scaled to 75-point system):")
        for cat in evaluation.category_results:
            # Calculate scaled score for this category
            cat_scaled = (cat.total_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            cat_max_scaled = (cat.max_score / evaluation.max_possible_raw_score * 75) if evaluation.max_possible_raw_score > 0 else 0
            
            print(f"   {cat.name}: {cat.total_score}/{cat.max_score} raw → {cat_scaled:.1f}/{cat_max_scaled:.1f} scaled ({cat.percentage:.1f}%)")
        
        print(f"\n📊 SCORING SUMMARY:")
        print(f"   • Content: {evaluation.scaled_content_score:.1f}/75 points (75% weight)")
        print(f"   • Attendance: {evaluation.attendance_score:.0f}/25 points (25% weight)")
        print(f"   • TOTAL: {evaluation.final_score:.1f}/100 points")
        
    except Exception as e:
        print(f"\n❌ ERROR during analysis: {str(e)}")
        print("\n💡 Troubleshooting tips:")
        print("   1. Make sure your .env file has valid API keys (ANTHROPIC_API_KEY or OPENAI_API_KEY)")
        print("   2. Check that the document path is correct and the file exists")
        print("   3. Ensure all required packages are installed:")
        print("      pip install python-docx openai anthropic python-dotenv")
        print("   4. The script will use mock analysis if API keys are not available")
        print("   5. Make sure you have write permissions to the output directory")


if __name__ == "__main__":
    main()
                