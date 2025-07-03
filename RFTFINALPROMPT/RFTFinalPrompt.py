import os
import time
import re
from datetime import datetime

from docx import Document
from openai import OpenAI

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ─── CONSTANT PROMPT TEMPLATE ────────────────────────────────────────────────────
FINAL_PROMPT = (
    """Flynn Construction Handover Analysis:\n\n{transcript}\n\n"
    "Analyze the above transcript and answer the question below.\n\n"
    "Please structure your answer using *exactly* these sections:\n"
    "1. Expert interpretation – concise statements explaining what was determined\n"
    "2. Supporting quotes – each on a new line, formatted as \"Speaker: 'exact quote'\"\n"
    "3. Professional summary – bullet list of any gaps, next steps, or items needing clarification\n\n"
    "Question: {question}\n"
    "Answer:"""
)


def read_transcript(transcript_path: str) -> str:
    """Read transcript from either .txt or .docx file."""
    file_extension = os.path.splitext(transcript_path)[1].lower()
    
    if file_extension == '.txt':
        # Read plain text file
        with open(transcript_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    
    elif file_extension == '.docx':
        # Read docx file
        doc = Document(transcript_path)
        # Extract all text from all paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text.strip())
        return '\n\n'.join(full_text)
    
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Only .txt and .docx files are supported.")


def _strip_markdown(text: str) -> str:
    """Remove **bold**, italics and other simple markdown artefacts."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text


def _tidy_labels(text: str) -> str:
    """Ensure headers appear once, with desired wording for RFT model."""
    replacements = {
        r"Synthesis": "Expert interpretation",
        r"Items\s*requiring\s*clarification": "Professional summary",
        r"Recommendations\s*/?\s*Items\s*requiring\s*clarification": "Professional summary",
    }
    lines = []
    for ln in text.splitlines():
        for pattern, repl in replacements.items():
            ln = re.sub(pattern, repl, ln, flags=re.IGNORECASE)
        lines.append(ln)
    # Deduplicate successive headers
    cleaned = []
    header_re = re.compile(r"^(?:\d+\.)?\s*(Expert interpretation|Supporting quotes|Professional summary)", re.I)
    prev = None
    for ln in lines:
        m = header_re.match(ln.strip())
        key = m.group(1).lower() if m else None
        if key and key == prev:
            continue
        cleaned.append(ln)
        if key:
            prev = key
    return "\n".join(cleaned)


def run_openai_rft_analysis(
    transcript_path: str,
    questions_path: str,
    output_path: str | None = None,
):
    """Run RFT analysis and export to a uniquely named Word (.docx) file."""
    # Load transcript (now supports both .txt and .docx)
    print(f"Loading transcript from: {transcript_path}")
    try:
        transcript = read_transcript(transcript_path)
        print(f"Successfully loaded transcript ({len(transcript)} characters)")
    except Exception as e:
        print(f"Error loading transcript: {e}")
        return
    
    # Load questions
    print(f"Loading questions from: {questions_path}")
    with open(questions_path, "r", encoding="utf-8") as f:
        questions = [line.strip() for line in f if line.strip()]
    print(f"Found {len(questions)} questions to process")
    
    # Determine versioned output path
    if output_path is None:
        output_dir = os.path.dirname(os.path.abspath(transcript_path))
        os.makedirs(output_dir, exist_ok=True)
        key = os.path.splitext(os.path.basename(transcript_path))[0]
        output_path = os.path.join(output_dir, f"rft_analysis_{key}.docx")
    print(f"\nReport will be saved to: {output_path}\n")
    
    # Create Word document
    doc = Document()
    project_name = os.path.splitext(os.path.basename(transcript_path))[0]
    doc.add_heading("FLYNN CONSTRUCTION – HANDOVER ANALYSIS", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(f"Model: ft:o4-mini-2025-04-16:flynn:test-eleven:BgYLe9DH")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    
    # Process each question
    for idx, question in enumerate(questions, 1):
        print(f"\rProcessing question {idx}/{len(questions)}…", end="", flush=True)
        prompt = FINAL_PROMPT.format(transcript=transcript, question=question)
        max_retries = 3
        for attempt in range(1, max_retries+1):
            try:
                resp = client.chat.completions.create(
                    model="ft:o4-mini-2025-04-16:flynn:test-eleven:BgYLe9DH",
                    messages=[{"role": "user", "content": prompt}],
                    max_completion_tokens=3000,
                )
                raw = resp.choices[0].message.content.strip()
                if len(raw) < 15:
                    raise ValueError("Answer too short, retrying...")
                clean = _tidy_labels(_strip_markdown(raw))
                doc.add_heading(f"Q{idx}: {question}", level=2)
                for para in clean.split("\n\n"):
                    doc.add_paragraph(para)
                break
            except Exception as e:
                if attempt < max_retries:
                    time.sleep(2)
                    continue
                doc.add_heading(f"Q{idx}: {question}", level=2)
                doc.add_paragraph(f"ERROR: Failed after {max_retries} attempts – {e}")
        time.sleep(0.5)
    
    # Save the Word document
    doc.save(output_path)
    print(f"\n✅ RFT analysis complete! Output saved to: {output_path}")


if __name__ == "__main__":
    # Example quick run - now supports .docx transcript
    run_openai_rft_analysis(
        transcript_path="C:\\Users\\kayla.dipaolo\\source\\repos\\provo river water treatment\\Provo River Water Treatment Turnover.docx",
        questions_path="C:\\Users\\kayla.dipaolo\\source\\repos\\base documents\\Handoff Questions.txt",
    )