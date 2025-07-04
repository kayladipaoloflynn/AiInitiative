#!/usr/bin/env python3
"""
Current Implementation Rubric Documentation Generator
Generates a Word document showing exactly what's implemented in the current analyzer code.

Requirements:
pip install python-docx

Usage:
Run this script to generate "Current_Rubric_Implementation.docx"
"""

import os
from datetime import datetime
from typing import Dict, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

class CurrentRubricDocumentationGenerator:
    """Generate documentation for the currently implemented rubric"""
    
    def __init__(self):
        self.doc = Document()
        self.rubric_data = self._load_current_rubric()
        
    def _load_current_rubric(self) -> Dict[str, Any]:
        """Load the exact rubric structure from current analyzer code"""
        return {
            "Customer Review": {
                "max_score": 10,
                "criteria": [
                    {
                        "name": "Customer Background Details",
                        "max_points": 10,
                        "description": "Customer organization and contacts named, relationship history documented, specific concerns identified",
                        "scoring_breakdown": {
                            "Customer Identity": {"points": 3, "guide": "3=Specific organization/contacts named | 2=Some names | 1=Generic identification | 0=Not identified"},
                            "Relationship History": {"points": 3, "guide": "3=Specific past projects documented | 2=Some history | 1=Basic mention | 0=No history"},
                            "Current Concerns": {"points": 3, "guide": "3=Specific requirements/concerns identified | 2=Some concerns | 1=Generic concerns | 0=No concerns"},
                            "Evidence Quality": {"points": 1, "guide": "1=Concrete details with proper nouns | 0=Generic statements"}
                        }
                    }
                ]
            },
            "Schedule Review": {
                "max_score": 15,
                "criteria": [
                    {
                        "name": "Project Timeline and Coordination",
                        "max_points": 15,
                        "description": "Specific start/end dates, crew planning, integration with other trades, milestone identification",
                        "scoring_breakdown": {
                            "Timeline Specificity": {"points": 4, "guide": "4=Specific start/end dates | 3=Clear dates | 2=General timeline | 1=Vague dates | 0=TBD"},
                            "Crew Planning": {"points": 4, "guide": "4=Crew size with man-day calculations | 3=Crew details | 2=Basic planning | 1=Minimal details | 0=No planning"},
                            "Trade Integration": {"points": 4, "guide": "4=Detailed coordination with other trades | 3=Good coordination | 2=Basic integration | 1=Limited coordination | 0=No coordination"},
                            "Milestone Management": {"points": 3, "guide": "3=Milestones with responsible parties | 2=Some milestones | 1=Basic milestones | 0=No milestones"}
                        }
                    }
                ]
            },
            "Material Review": {
                "max_score": 15,
                "criteria": [
                    {
                        "name": "Material Planning and Logistics",
                        "max_points": 15,
                        "description": "Specific materials with quantities, suppliers identified, lead times, installation sequence",
                        "scoring_breakdown": {
                            "Material Specificity": {"points": 4, "guide": "4=Specific materials with quantities | 3=Good details | 2=Some specifics | 1=Basic list | 0=Generic materials"},
                            "Supplier Information": {"points": 4, "guide": "4=Suppliers identified by name | 3=Some supplier info | 2=Basic supplier data | 1=Limited info | 0=No supplier info"},
                            "Lead Time Planning": {"points": 4, "guide": "4=Specific lead times with dates | 3=Good timing info | 2=Basic timing | 1=Limited timing | 0=No lead times"},
                            "Installation Process": {"points": 3, "guide": "3=Detailed installation sequence | 2=Some process info | 1=Basic process | 0=No installation planning"}
                        }
                    }
                ]
            },
            "Risk Review": {
                "max_score": 15,
                "criteria": [
                    {
                        "name": "Risk Assessment and Mitigation",
                        "max_points": 15,
                        "description": "Project-specific risks identified, detailed mitigation strategies, site-specific hazard analysis",
                        "scoring_breakdown": {
                            "Risk Identification": {"points": 5, "guide": "5=Project-specific risks (not generic) | 4=Good risk ID | 3=Some specific risks | 2=Mix of generic/specific | 1=Mostly generic | 0=Only generic"},
                            "Mitigation Strategies": {"points": 5, "guide": "5=Detailed mitigation for each risk | 4=Good strategies | 3=Some mitigation | 2=Basic strategies | 1=Limited mitigation | 0=No mitigation"},
                            "Site-Specific Analysis": {"points": 3, "guide": "3=Comprehensive site analysis | 2=Good site considerations | 1=Basic site factors | 0=No site-specific analysis"},
                            "Prevention Methods": {"points": 2, "guide": "2=Specific prevention methods described | 1=Some prevention | 0=No prevention methods"}
                        }
                    }
                ]
            },
            "Equipment Review": {
                "max_score": 10,
                "criteria": [
                    {
                        "name": "Equipment Planning and Coordination",
                        "max_points": 10,
                        "description": "Specific equipment types, costs/rates, provider identification, delivery timeline, responsibility assignments",
                        "scoring_breakdown": {
                            "Equipment Specificity": {"points": 3, "guide": "3=Specific equipment types listed | 2=Good equipment details | 1=Basic equipment info | 0=Generic equipment"},
                            "Cost Information": {"points": 2, "guide": "2=Costs or rental rates provided | 1=Some cost info | 0=No cost information"},
                            "Provider Details": {"points": 2, "guide": "2=Equipment provider identified | 1=Some provider info | 0=No provider information"},
                            "Timeline and Responsibility": {"points": 2, "guide": "2=Delivery timeline and clear responsibilities | 1=Some timing/responsibility info | 0=No timeline/responsibility"},
                            "Analysis Quality": {"points": 1, "guide": "1=Thorough analysis vs 'standard equipment' | 0=Generic responses"}
                        }
                    }
                ]
            },
            "Project Scope Review": {
                "max_score": 20,
                "criteria": [
                    {
                        "name": "Drawing Review and Scope Definition",
                        "max_points": 20,
                        "description": "Evidence of actual drawing review, scope clearly defined, estimating handoff complete, technical documentation",
                        "scoring_breakdown": {
                            "Drawing Review Evidence": {"points": 6, "guide": "6=Specific drawing numbers/dates referenced | 5=Good drawing evidence | 4=Some drawing specifics | 3=Basic drawing review | 2=Limited evidence | 1=Minimal evidence | 0=No evidence"},
                            "Scope Definition": {"points": 5, "guide": "5=Clear scope with inclusions/exclusions | 4=Well-defined scope | 3=Generally clear scope | 2=Some scope clarity | 1=Basic scope | 0=Unclear scope"},
                            "Estimating Handoff": {"points": 5, "guide": "5=Complete estimating handoff documented | 4=Good handoff | 3=Adequate handoff | 2=Some handoff | 1=Basic handoff | 0=No handoff"},
                            "Technical Documentation": {"points": 3, "guide": "3=Technical docs provided (Bluebeam, markups) | 2=Some documentation | 1=Basic docs | 0=No documentation"},
                            "Response Quality": {"points": 1, "guide": "1=Project-specific details vs templates | 0=Templated responses"}
                        }
                    }
                ]
            },
            "Foreman Comments": {
                "max_score": 10,
                "criteria": [
                    {
                        "name": "Foreman Input and Participation",
                        "max_points": 10,
                        "description": "Foreman present and active, specific suggestions/concerns, actionable items, expertise applied",
                        "scoring_breakdown": {
                            "Participation Level": {"points": 4, "guide": "4=Active participation with expertise | 3=Good participation | 2=Some participation | 1=Limited participation | 0=No participation"},
                            "Specific Input": {"points": 3, "guide": "3=Specific suggestions/concerns documented | 2=Some specific input | 1=Basic input | 0=No specific input"},
                            "Actionable Items": {"points": 2, "guide": "2=Clear actionable items identified | 1=Some actionable items | 0=No actionable items"},
                            "Evidence of Expertise": {"points": 1, "guide": "1=Clear evidence of foreman expertise applied | 0=No evidence of expertise"}
                        }
                    }
                ]
            },
            "Payroll Review": {
                "max_score": 5,
                "criteria": [
                    {
                        "name": "Payroll Setup and Verification",
                        "max_points": 5,
                        "description": "Specific payroll attributes confirmed, labor codes verified, setup completion documented",
                        "scoring_breakdown": {
                            "Payroll Attributes": {"points": 2, "guide": "2=Specific payroll attributes confirmed | 1=Some attributes | 0=Generic review"},
                            "Labor Codes": {"points": 2, "guide": "2=Labor codes and classifications verified | 1=Some verification | 0=No verification"},
                            "Setup Completion": {"points": 1, "guide": "1=Setup completion documented | 0=Not confirmed"}
                        }
                    }
                ]
            }
        }
    
    def generate_current_rubric_documentation(self, output_path: str = "C:\\Users\\kayla.dipaolo\\source\\repos\\rubric analyzer code\\rubric 3.0.docx"):
        """Generate documentation matching the current implementation"""
        
        self._setup_styles()
        self._add_title_page()
        self._add_overview()
        self._add_scoring_methodology()
        self._add_detailed_sections()
        self._add_performance_levels()
        self._add_foreman_assessment()
        self._add_implementation_notes()
        
        self.doc.save(output_path)
        print(f"📋 Current rubric documentation saved to: {output_path}")
    
    def _setup_styles(self):
        """Setup professional document styles"""
        styles = self.doc.styles
        
        # Title style
        try:
            if 'Rubric Title' not in [s.name for s in styles]:
                title_style = styles.add_style('Rubric Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.name = 'Calibri'
                title_style.font.size = Pt(24)
                title_style.font.bold = True
                title_style.font.color.rgb = RGBColor(0, 51, 102)
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
        
        # Section Header style
        try:
            if 'Section Header' not in [s.name for s in styles]:
                section_style = styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
                section_style.font.name = 'Calibri'
                section_style.font.size = Pt(16)
                section_style.font.bold = True
                section_style.font.color.rgb = RGBColor(0, 102, 204)
        except:
            pass
    
    def _add_title_page(self):
        """Add title page"""
        
        # Main title
        try:
            title = self.doc.add_paragraph("Project Handoff Evaluation Rubric", style='Rubric Title')
        except:
            title = self.doc.add_paragraph("Project Handoff Evaluation Rubric")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Subtitle
        subtitle = self.doc.add_paragraph("Current Implementation - Specificity Focused")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Version info
        version_para = self.doc.add_paragraph()
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_para.add_run("Version: ").bold = True
        version_para.add_run("2.0 - Specificity Based")
        
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run("Generated: ").bold = True
        date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        
        self.doc.add_page_break()
    
    def _add_overview(self):
        """Add overview section"""
        
        self.doc.add_heading('Overview', level=1)
        
        overview_text = """
This rubric evaluates project handoff documentation and meetings using an 8-section framework 
focused on content specificity and quality. The evaluation prioritizes concrete, project-specific 
details over generic responses.

The system uses AI analysis to distinguish between high-quality, specific content and 
generic, templated responses. This approach ensures handoffs demonstrate actual preparation 
and project understanding rather than form completion.
"""
        
        self.doc.add_paragraph(overview_text.strip())
        
        # Key principles
        self.doc.add_heading('Evaluation Principles', level=2)
        
        principles = [
            "Reward specific details: proper nouns, exact dates, quantities, dollar amounts",
            "Penalize generic responses: 'standard materials', 'will coordinate', 'TBD'",
            "Focus on project-specific information rather than templated language",
            "Evidence of actual preparation vs. form completion",
            "Foreman participation with meaningful, technical input"
        ]
        
        for principle in principles:
            p = self.doc.add_paragraph(f"• {principle}")
    
    def _add_scoring_methodology(self):
        """Add scoring methodology section"""
        
        self.doc.add_heading('Scoring Methodology', level=1)
        
        # Overall approach
        self.doc.add_heading('Weighted Scoring System', level=2)
        methodology_text = """
The rubric uses a scaled 100-point system combining content quality assessment 
with foreman attendance verification:

• Content Quality (75 points): Evaluation across 8 sections totaling 100 raw points
• Foreman Attendance (25 points): Binary assessment based on meaningful participation

Final Score = (Content Score ÷ 100 × 75) + (Attendance Score)
"""
        self.doc.add_paragraph(methodology_text.strip())
        
        # Content breakdown
        self.doc.add_heading('Content Sections (100 Raw Points → 75 Scaled Points)', level=2)
        
        content_table = self.doc.add_table(rows=1, cols=3)
        try:
            content_table.style = 'Table Grid'
        except:
            pass
        
        # Headers
        headers = ['Section', 'Raw Points', 'Description']
        header_cells = content_table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            try:
                header_cells[i].paragraphs[0].runs[0].bold = True
            except:
                pass
        
        # Add content sections
        for section_name, section_data in self.rubric_data.items():
            row_cells = content_table.add_row().cells
            row_cells[0].text = section_name
            row_cells[1].text = f"{section_data['max_score']} points"
            row_cells[2].text = section_data['criteria'][0]['description']
    
    def _add_detailed_sections(self):
        """Add detailed section breakdown"""
        
        self.doc.add_heading('Detailed Section Breakdown', level=1)
        
        section_number = 1
        
        for section_name, section_data in self.rubric_data.items():
            # Section header
            try:
                section_header = self.doc.add_paragraph(
                    f"Section {section_number}: {section_name} ({section_data['max_score']} points)", 
                    style='Section Header'
                )
            except:
                section_header = self.doc.add_paragraph(f"Section {section_number}: {section_name} ({section_data['max_score']} points)")
                try:
                    section_header.style = self.doc.styles['Heading 2']
                except:
                    pass
            
            # Section description
            criterion = section_data['criteria'][0]
            desc_para = self.doc.add_paragraph()
            desc_para.add_run("Description: ").bold = True
            desc_para.add_run(criterion['description'])
            
            # Scoring breakdown
            self.doc.add_heading('Scoring Components:', level=3)
            
            breakdown_table = self.doc.add_table(rows=1, cols=3)
            try:
                breakdown_table.style = 'Table Grid'
            except:
                pass
            
            # Table headers
            breakdown_headers = ['Component', 'Points', 'Scoring Guide']
            header_cells = breakdown_table.rows[0].cells
            for i, header in enumerate(breakdown_headers):
                header_cells[i].text = header
                try:
                    header_cells[i].paragraphs[0].runs[0].bold = True
                except:
                    pass
            
            # Add breakdown rows
            for component, details in criterion['scoring_breakdown'].items():
                row_cells = breakdown_table.add_row().cells
                row_cells[0].text = component
                row_cells[1].text = str(details['points'])
                row_cells[2].text = details['guide']
            
            section_number += 1
            self.doc.add_paragraph()  # Add spacing
    
    def _add_performance_levels(self):
        """Add performance level definitions"""
        
        self.doc.add_heading('Performance Level Classifications', level=1)
        
        perf_table = self.doc.add_table(rows=6, cols=3)
        try:
            perf_table.style = 'Table Grid'
        except:
            pass
        
        # Headers
        headers = ['Performance Level', 'Score Range', 'Characteristics']
        header_cells = perf_table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            try:
                header_cells[i].paragraphs[0].runs[0].bold = True
            except:
                pass
        
        # Performance levels
        levels = [
            ("EXCELLENT", "60-100 points", "High specificity, comprehensive project details, strong foreman input"),
            ("GOOD", "40-59 points", "Good specificity with some details, adequate preparation evident"),
            ("SATISFACTORY", "20-39 points", "Basic specificity, some gaps in detail or preparation"),
            ("NEEDS IMPROVEMENT", "10-19 points", "Limited specificity, significant preparation gaps"),
            ("UNSATISFACTORY", "0-9 points", "Generic responses, poor preparation, minimal specifics")
        ]
        
        for level, range_val, desc in levels:
            row_cells = perf_table.add_row().cells
            row_cells[0].text = level
            row_cells[1].text = range_val
            row_cells[2].text = desc
    
    def _add_foreman_assessment(self):
        """Add foreman attendance assessment details"""
        
        self.doc.add_heading('Foreman Attendance Assessment (25 Points)', level=1)
        
        foreman_text = """
The foreman attendance component is evaluated separately from content quality and contributes 
25% of the total score. This assessment uses AI analysis to detect foreman presence and 
meaningful participation.
"""
        self.doc.add_paragraph(foreman_text.strip())
        
        # Scoring criteria
        self.doc.add_heading('Scoring Criteria', level=2)
        
        criteria_table = self.doc.add_table(rows=3, cols=2)
        try:
            criteria_table.style = 'Table Grid'
        except:
            pass
        
        # Headers
        criteria_headers = ['Score', 'Requirements']
        header_cells = criteria_table.rows[0].cells
        for i, header in enumerate(criteria_headers):
            header_cells[i].text = header
            try:
                header_cells[i].paragraphs[0].runs[0].bold = True
            except:
                pass
        
        # Scoring data
        scoring_data = [
            ("25 points", "Foreman present AND provides meaningful, project-specific input AND contributes actionable suggestions"),
            ("0 points", "Foreman not present OR no meaningful input OR only generic responses")
        ]
        
        for score, requirement in scoring_data:
            row_cells = criteria_table.add_row().cells
            row_cells[0].text = score
            row_cells[1].text = requirement
        
        # Detection method
        self.doc.add_heading('AI Detection Method', level=2)
        detection_text = """
The system analyzes the handoff documentation for evidence of foreman participation through:

• Direct quotes or paraphrased foreman input
• Specific technical suggestions or field expertise
• Project-specific concerns raised by field personnel
• Actionable items attributed to foreman
• Discussion of constructability or installation challenges

The AI focuses on ROLE and CONTRIBUTIONS rather than job titles, identifying field supervisors 
who provide installation insights and practical input.
"""
        self.doc.add_paragraph(detection_text.strip())
    
    def _add_implementation_notes(self):
        """Add implementation and usage notes"""
        
        self.doc.add_heading('Implementation Notes', level=1)
        
        # AI analysis approach
        self.doc.add_heading('AI Analysis Approach', level=2)
        ai_text = """
The evaluation uses advanced language models (Claude 3.5 Sonnet or GPT-4) to analyze 
each section against specific criteria. The AI is prompted to:

• Reward concrete, project-specific details
• Penalize generic or templated responses
• Focus on evidence of actual preparation
• Distinguish between quality content and form completion
• Provide detailed justification for scoring decisions
"""
        self.doc.add_paragraph(ai_text.strip())
        
        # Quality indicators
        self.doc.add_heading('Quality Indicators', level=2)
        
        self.doc.add_heading('High-Quality Response Indicators:', level=3)
        high_quality = [
            "Proper nouns (company names, contact names, locations)",
            "Specific dates and quantities",
            "Dollar amounts and time estimates",
            "Evidence of coordination planning",
            "Project-specific risk assessments",
            "Actionable next steps with responsibilities"
        ]
        
        for indicator in high_quality:
            p = self.doc.add_paragraph(f"• {indicator}")
        
        self.doc.add_heading('Low-Quality Response Indicators:', level=3)
        low_quality = [
            "'We will coordinate' without specifics",
            "'TBD' or 'to be determined'",
            "'Standard' or 'typical' without details",
            "Single-sentence responses to complex topics",
            "Generic safety checklists without project context",
            "No evidence of foreman engagement"
        ]
        
        for indicator in low_quality:
            p = self.doc.add_paragraph(f"• {indicator}")
        
        # File format support
        self.doc.add_heading('Supported File Formats', level=2)
        format_text = """
The analyzer supports multiple input formats:

• Word documents (.docx)
• JSON files (converted to readable text)
• CSV files (processed as tabular data)
• Plain text files (.txt)

The system automatically detects file type and converts content to analyzable text format.
"""
        self.doc.add_paragraph(format_text.strip())


def main():
    """Generate the current rubric documentation"""
    
    print("=" * 60)
    print("CURRENT RUBRIC IMPLEMENTATION DOCUMENTATION GENERATOR")
    print("=" * 60)
    
    try:
        generator = CurrentRubricDocumentationGenerator()
        
        # Generate the documentation
        output_path = "Current_Rubric_Implementation.docx"
        generator.generate_current_rubric_documentation(output_path)
        
        # Print summary
        total_sections = len(generator.rubric_data)
        total_points = sum(section_data['max_score'] for section_data in generator.rubric_data.values())
        
        print(f"✅ Documentation generated successfully!")
        print(f"📊 Total Sections: {total_sections}")
        print(f"🎯 Total Raw Points: {total_points} (scaled to 75)")
        print(f"👥 Foreman Attendance: 25 points")
        print(f"📁 Saved to: {output_path}")
        print("\nThe document includes:")
        print("  • Complete 8-section rubric breakdown")
        print("  • Detailed scoring guidelines for each component")
        print("  • Performance level definitions")
        print("  • Foreman attendance assessment method")
        print("  • AI analysis approach and quality indicators")
        print("  • Implementation notes and file format support")
        
    except Exception as e:
        print(f"❌ Error generating documentation: {e}")
        print("\nTroubleshooting:")
        print("  • Ensure python-docx is installed: pip install python-docx")
        print("  • Check that you have write permissions in the current directory")


if __name__ == "__main__":
    main()