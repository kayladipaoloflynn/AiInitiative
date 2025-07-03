import os
import time
import re
from datetime import datetime

from docx import Document
from openai import OpenAI

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ─── CONSTANT PROMPT TEMPLATE ────────────────────────────────────────────────────
FINAL_PROMPT = (
    """Flynn Construction Handover Analysis:\n\n{transcript}\n\n"
    "Analyze the above transcript and answer the question below.\n\n"
    "Please structure your answer using *exactly* these sections:\n"
    "1. Expert interpretation – concise statements explaining what was determined\n"
    "2. Supporting quotes – each on a new line, formatted as \"Speaker: 'exact quote'\"\n"
    "3. Professional summary – bullet list of any gaps, next steps, or items needing clarification\n\n"
    "Question: {question}\n"
    "Answer:"""
)


def read_transcript(transcript_path: str) -> str:
    """Read transcript from either .txt or .docx file."""
    file_extension = os.path.splitext(transcript_path)[1].lower()
    
    if file_extension == '.txt':
        # Read plain text file
        with open(transcript_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    
    elif file_extension == '.docx':
        # Read docx file
        doc = Document(transcript_path)
        # Extract all text from all paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text.strip())
        return '\n\n'.join(full_text)
    
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Only .txt and .docx files are supported.")


def run_openai_analysis(
    transcript_path: str,
    questions_path: str,
    output_path: str | None = None,
):
    """Generate a .docx analysis, versioned by transcript name to avoid overwrite."""
    # Load transcript (now supports both .txt and .docx)
    print(f"Loading transcript from: {transcript_path}")
    try:
        transcript = read_transcript(transcript_path)
        print(f"Successfully loaded transcript ({len(transcript)} characters)")
    except Exception as e:
        print(f"Error loading transcript: {e}")
        return
    
    # Load questions
    print(f"Loading questions from: {questions_path}")
    with open(questions_path, "r", encoding="utf-8") as f:
        questions = [line.strip() for line in f if line.strip()]
    print(f"Found {len(questions)} questions to process")
    
    # Determine output path
    if output_path is None:
        # folder next to transcript
        output_dir = os.path.dirname(os.path.abspath(transcript_path))
        os.makedirs(output_dir, exist_ok=True)
        # derive unique name per transcript
        project_key = os.path.splitext(os.path.basename(transcript_path))[0]
        filename = f"dpo_analysis_{project_key}.docx"
        output_path = os.path.join(output_dir, filename)
    
    print("\nStarting analysis …")
    print(f"Word report will be saved to: {output_path}")
    
    # Create document
    doc = Document()
    project_name = os.path.splitext(os.path.basename(transcript_path))[0]
    doc.add_heading("FLYNN CONSTRUCTION – HANDOVER ANALYSIS", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(f"Model: ft:gpt-4o-2024-08-06:flynn:test-nine:BcuSbEPv")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Process each question
    for idx, question in enumerate(questions, 1):
        print(f"\rProcessing question {idx}/{len(questions)}…", end="", flush=True)
        prompt = FINAL_PROMPT.format(transcript=transcript, question=question)
        try:
            response = client.chat.completions.create(
                model="ft:gpt-4o-2024-08-06:flynn:test-nine:BcuSbEPv",
                messages=[{"role": "user", "content": prompt}],
                max_completion_tokens=2500,
            )
            answer = response.choices[0].message.content.strip()
            # Tidy markdown
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", answer)
            # Write Q&A
            doc.add_heading(f"Q{idx}: {question}", level=2)
            for block in clean.split("\n\n"):
                doc.add_paragraph(block)
            doc.add_paragraph()
        except Exception as e:
            print(f"\nError on question {idx}: {e}")
            doc.add_heading(f"Q{idx}: {question}", level=2)
            doc.add_paragraph(f"ERROR: {e}")
            doc.add_paragraph()
        time.sleep(0.5)
    
    # Save document
    doc.save(output_path)
    print(f"\n\n✅ Analysis complete! Output saved to: {output_path}")


def quick_run():
    # default paths - now supports .docx transcript
    transcript_path = (
        "C:\\Users\\kayla.dipaolo\\source\\repos\\provo river water treatment\\Provo River Water Treatment Turnover.docx"
    )
    questions_path = (
        "C:\\Users\\kayla.dipaolo\\source\\repos\\base documents\\Handoff Questions.txt"
    )
    run_openai_analysis(transcript_path, questions_path)


if __name__ == "__main__":
    quick_run()