import anthropic
import os
import time
from datetime import datetime
from docx import Document

# Initialize the Anthropic client
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

# ─── CONFIGURATION ─────────────────────────────────────────────────────────────
MODEL = "claude-3-5-sonnet-20241022"  # or "claude-3-opus-20240229"
TEMPERATURE = 0.2
MAX_TOKENS = 2000

# ─── PROMPT TEMPLATES ───────────────────────────────────────────────────────────
SYSTEM_PROMPT = (
    "You are an expert transcript analyst for Flynn Construction, specializing in "
    "construction project handover meetings.\n\nKey principles:\n"
    "- Provide professional-level analysis (assume readers understand construction)\n"
    "- Quote directly from the transcript to support all claims\n"
    "- Synthesize information clearly before presenting evidence\n"
    "- Only suggest follow-ups for genuinely unclear contract/spec items\n"
    "- Keep responses concise but comprehensive"
)

USER_PROMPT_TEMPLATE = (
    "You are a senior project manager at Flynn Construction preparing your team for this project.\n\n"
    "{transcript}\n\n"
    "Analyze the handover meeting and provide actionable information.\n\n"
    "Please structure your answer using *exactly* these sections:\n"
    "1. Expert interpretation – concise statements explaining what was determined\n"
    "2. Supporting quotes – each on a new line, formatted as \"Speaker: 'exact quote'\"\n"
    "3. Professional summary – bullet list of any gaps, next steps, or items needing clarification\n\n"
    "Focus on what the construction team needs to execute successfully.\n\n"
    "Question: {question}\nAnswer:"
)

# ─── FILE LOADING FUNCTIONS ─────────────────────────────────────────────────────

def load_transcript(file_path: str) -> str:
    """
    Load transcript from either a .txt or .docx file.
    
    Args:
        file_path: Path to the transcript file (.txt or .docx)
        
    Returns:
        str: The transcript content
        
    Raises:
        ValueError: If file extension is not supported
        FileNotFoundError: If file doesn't exist
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Transcript file not found: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.txt':
        print(f"Loading .txt transcript from: {file_path}")
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    
    elif file_ext == '.docx':
        print(f"Loading .docx transcript from: {file_path}")
        doc = Document(file_path)
        # Extract all text from all paragraphs
        transcript_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                transcript_parts.append(paragraph.text.strip())
        return "\n\n".join(transcript_parts)
    
    else:
        raise ValueError(f"Unsupported file format: {file_ext}. Only .txt and .docx files are supported.")

# ─── MAIN ANALYSIS FUNCTION ─────────────────────────────────────────────────────

def run_anthropic_analysis(transcript_path: str, questions_path: str, output_path: str | None = None):
    """
    Run the Anthropic analysis and write answers into a Word document (.docx).
    Each transcript generates a uniquely named output file to prevent overwrites.
    
    Args:
        transcript_path: Path to transcript file (.txt or .docx)
        questions_path: Path to questions file (.txt)
        output_path: Optional custom output path for the analysis document
    """
    # Load transcript (now supports both .txt and .docx)
    try:
        transcript = load_transcript(transcript_path)
    except Exception as e:
        print(f"Error loading transcript: {e}")
        return

    # Load questions
    print(f"Loading questions from: {questions_path}")
    try:
        with open(questions_path, "r", encoding="utf-8") as f:
            questions = [line.strip() for line in f if line.strip()]
    except FileNotFoundError:
        print(f"Questions file not found: {questions_path}")
        return
    except Exception as e:
        print(f"Error loading questions: {e}")
        return

    print(f"Found {len(questions)} questions to process")

    # Determine output path (next to transcript, unique per file)
    if output_path is None:
        output_dir = os.path.dirname(os.path.abspath(transcript_path))
        os.makedirs(output_dir, exist_ok=True)
        project_key = os.path.splitext(os.path.basename(transcript_path))[0]
        output_path = os.path.join(output_dir, f"anthropic_analysis_{project_key}.docx")

    print("\nStarting analysis …")
    print(f"Word report will be saved to: {output_path}\n")

    # Create Word document
    doc = Document()
    project_name = os.path.splitext(os.path.basename(transcript_path))[0]
    doc.add_heading("FLYNN CONSTRUCTION – HANDOVER ANALYSIS", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(f"Model: {MODEL}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Transcript Format: {os.path.splitext(transcript_path)[1].upper()}")

    # Process each question
    for idx, question in enumerate(questions, 1):
        print(f"\rProcessing {idx}/{len(questions)}…", end="", flush=True)

        user_prompt = USER_PROMPT_TEMPLATE.format(transcript=transcript, question=question)

        try:
            response = client.messages.create(
                model=MODEL,
                system=SYSTEM_PROMPT,
                max_tokens=MAX_TOKENS,
                temperature=TEMPERATURE,
                messages=[{"role": "user", "content": user_prompt}],
            )
            answer = response.content[0].text.strip()

            # Write Q&A into the document
            doc.add_heading(f"Q{idx}: {question}", level=2)
            for para in answer.split("\n\n"):
                doc.add_paragraph(para)
            doc.add_paragraph()  # blank line for spacing

        except Exception as e:
            print(f"\nError on question {idx}: {e}")
            doc.add_heading(f"Q{idx}: {question}", level=2)
            doc.add_paragraph(f"ERROR: {e}")
            doc.add_paragraph()

        time.sleep(0.5)  # courtesy delay for rate‑limiting

    # Save the Word document
    try:
        doc.save(output_path)
        print(f"\n\n✅ Analysis complete! Output saved to: {output_path}")
    except Exception as e:
        print(f"\n\n❌ Error saving document: {e}")

# ─── QUICK‑RUN HELPER ───────────────────────────────────────────────────────────

def quick_run():
    transcript_path = (
        "C:\\Users\\kayla.dipaolo\\source\\repos\\provo river water treatment\\Provo River Water Treatment Turnover.docx"
    )
    questions_path = (
        "C:\\Users\\kayla.dipaolo\\source\\repos\\base documents\\Handoff Questions.txt"
    )
    run_anthropic_analysis(transcript_path, questions_path)

if __name__ == "__main__":
    quick_run()